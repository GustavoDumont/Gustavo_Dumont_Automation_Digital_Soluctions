"""
I LOVE Cmat — App em CustomTkinter
Design: Dark / Azul escuro

Funcionalidades (em massa, sem limite explícito de arquivos):
 - Colocar marca d'água (PDF watermark — texto ou PDF como selo)
 - Juntar (merge) vários PDFs
 - Organizar PDFs (reordenar páginas por padrão, dividir por intervalos, renomear em lote)
 - Converter imagens (JPG/PNG/...) para PDF (cada imagem vira 1 PDF; opcional: juntar todas em um único PDF)
 - Converter .docx para PDF (usa docx2pdf no Windows; tenta soffice/LibreOffice em outros sistemas)

Instalação (recomendada):
 pip install customtkinter pypdf pillow reportlab docx2pdf

Observações:
 - docx2pdf funciona melhor no Windows com MS Word instalado. Em outros sistemas, o app tentará chamar o `soffice` (LibreOffice) se disponível.
 - O app usa threads para não travar a interface e mostra barra de progresso simples.
 - Use com permissões de leitura/escrita nas pastas alvo.
"""

import os
import sys
import threading
import tempfile
import subprocess
import io
import shutil
import fitz
import tkinter as tk
import re
from pathlib import Path
from typing import List, Tuple
from io import BytesIO

import customtkinter as ctk
from tkinter import filedialog, messagebox, simpledialog
from pypdf import PdfReader, PdfWriter
from PIL import Image, ImageTk
from pdf2image import convert_from_path
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib.utils import ImageReader
from reportlab.lib import colors

# Try optional import for Word->PDF convenience
try:
    from docx2pdf import convert as docx2pdf_convert
except Exception:
    docx2pdf_convert = None

ctk.set_appearance_mode("Dark")  # Dark mode
ctk.set_default_color_theme("blue")  # azul escuro estilo

APP_NAME = "I LOVE Cmat"

#============================= Definições para o Renomeador de Arquivos ========================
# -------------------------------------------------------
# Tesseract portátil e fuzzy search
# -------------------------------------------------------
import sys, os
try:
    # base do app (modo portátil com PyInstaller)
    if getattr(sys, "frozen", False):
        BASE_DIR = sys._MEIPASS
    else:
        BASE_DIR = os.path.dirname(os.path.abspath(__file__))
except Exception:
    BASE_DIR = os.path.dirname(os.path.abspath(__file__))

TESSERACT_PORTABLE = os.path.join(BASE_DIR, "tesseract", "tesseract.exe")
TESSDATA_PORTABLE = os.path.join(BASE_DIR, "tesseract", "tessdata")

try:
    import pytesseract
    if os.path.exists(TESSERACT_PORTABLE):
        pytesseract.pytesseract.tesseract_cmd = TESSERACT_PORTABLE
        os.environ["TESSDATA_PREFIX"] = TESSDATA_PORTABLE
    # else assume user has tesseract in PATH
except Exception as e:
    print("pytesseract não disponível:", e)

# fuzzy lib
try:
    from rapidfuzz import fuzz
    USE_RAPIDFUZZ = True
except Exception:
    import difflib
    USE_RAPIDFUZZ = False

import importlib.util
import subprocess
import sys

#============================= Definições para Converter Word em PDF ========================

# Garante que 'python-docx' está instalado
if importlib.util.find_spec("docx") is None:
    try:
        subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    except Exception as e:
        print("Falha ao instalar python-docx:", e)

############################## Funções Auxiliares ##############################################

# ------------------- Auxiliares para Organizar PDFs -------------------
# Estado global da pré-visualização (declarar uma única vez no topo do arquivo)
preview_state = {"render_thread": None, "cancel_flag": False, "current_page": None}

import threading
import fitz  # PyMuPDF
from PIL import Image, ImageTk

# --- Nova lógica de clique e arrasto ---
drag_state = {"start_x": 0, "start_y": 0, "widget": None, "index": None, "moved": False}

def on_thumb_press(event, page):
    widget = event.widget
    drag_state.update({
        "start_x": event.x_root,
        "start_y": event.y_root,
        "widget": widget,
        "index": thumb_widgets.index(widget),
        "moved": False
    })
    widget.config(relief="ridge")


def on_thumb_motion(event):
    if not drag_state["widget"]:
        return

    dx = abs(event.x_root - drag_state["start_x"])
    dy = abs(event.y_root - drag_state["start_y"])
    if dx > 10 or dy > 10:  # arrasto detectado
        drag_state["moved"] = True


def on_thumb_release(event):
    widget = drag_state["widget"]
    if not widget:
        return

    widget.config(relief="flat")

    # Se não houve movimento → clique simples (mostrar preview)
    if not drag_state["moved"]:
        try:
            page = getattr(widget, "page_num", None)
            file_path = getattr(widget, "file_path", None)

            if page is not None and file_path:
                print(f"[CLICK DETECTADO] Miniatura {page}")
                safe_update_preview(file_path, page, preview_canvas)
            else:
                print("[ERRO] Miniatura sem atributos necessários (page_num ou file_path ausentes)")
        except Exception as e:
            print(f"[click err] {e}")

        drag_state.update({"widget": None, "index": None, "moved": False})
        return

    # Se arrastou → reorganizar
    try:
        x = event.x_root - inner_frame.winfo_rootx()
        y = event.y_root - inner_frame.winfo_rooty()
        dest_col = max(0, int(x // 280))
        dest_row = max(0, int(y // 360))
        dest_index = min(dest_row * 3 + dest_col, len(thumb_widgets) - 1)

        src_index = drag_state["index"]
        item = thumb_widgets.pop(src_index)
        thumb_widgets.insert(dest_index, item)

        # Reposiciona tudo
        for idx, w in enumerate(thumb_widgets):
            w.grid(row=idx // 3, column=idx % 3, padx=8, pady=8)

        print(f"[MOVIDO] Página {item.page_num} para posição {dest_index + 1}")
    except Exception as e:
        print(f"[move err] {e}")

    drag_state.update({"widget": None, "index": None, "moved": False})

# ---------------- FUNÇÕES DE PRÉ-VISUALIZAÇÃO ----------------
def safe_update_preview(pdf_path, page, preview_canvas):
    """
    Atualiza o preview de forma segura (mostra apenas a página clicada).
    """
    global preview_state

    # Inicializa o estado se ainda não existir
    if "render_thread" not in preview_state:
        preview_state.update({
            "render_thread": None,
            "cancel_flag": False,
            "current_page": None
        })

    # Cancela renderização anterior se estiver ativa
    if preview_state["render_thread"] and preview_state["render_thread"].is_alive():
        preview_state["cancel_flag"] = True
        print("[debug] Cancelando renderização anterior...")

    # Atualiza o estado
    preview_state["cancel_flag"] = False
    preview_state["current_page"] = page

    # Limpa e mostra texto temporário
    preview_canvas.delete("all")
    preview_canvas.create_text(
        240, 320,
        text=f"Carregando página {page}...",
        fill="gray",
        font=("Arial", 12, "italic")
    )
    preview_canvas.update_idletasks()

    # --- Função interna que faz o trabalho pesado ---
    def render_page():
        try:
            doc = fitz.open(pdf_path)
            if preview_state["cancel_flag"]:
                doc.close()
                return

            # Renderiza a página
            page_obj = doc.load_page(page - 1)
            mat = fitz.Matrix(1.5, 1.5)
            pix = page_obj.get_pixmap(matrix=mat)
            doc.close()

            if preview_state["cancel_flag"]:
                return

            pil_img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)

            # Ajusta ao tamanho atual do canvas
            cw, ch = preview_canvas.winfo_width(), preview_canvas.winfo_height()
            if cw <= 1 or ch <= 1:
                cw, ch = 480, 650  # fallback seguro

            iw, ih = pil_img.size
            scale = min(cw / iw, ch / ih)
            pil_img = pil_img.resize((int(iw * scale), int(ih * scale)), Image.LANCZOS)
            tk_img = ImageTk.PhotoImage(pil_img)

            # Atualiza o canvas no thread principal
            def update_canvas():
                if preview_state["cancel_flag"]:
                    return
                preview_canvas.delete("all")
                x_center, y_center = cw // 2, ch // 2
                preview_canvas.create_image(x_center, y_center, image=tk_img, anchor="center")
                preview_canvas.image = tk_img
                preview_canvas.full_image = pil_img
                print(f"[debug] -> Exibindo página {page}")

            preview_canvas.after(0, update_canvas)

        except Exception as e:
            print(f"[preview err] {e}")

    # Cria a nova thread de renderização
    t = threading.Thread(target=render_page, daemon=True)
    preview_state["render_thread"] = t
    t.start()
        
def update_preview(pdf_path, page_num, preview_canvas, preview_state, zoom=1.5):
    def _render():
        try:
            # Evita redesenhar a mesma página
            if preview_state.get("current_page") == page_num:
                return

            print(f"[debug] Atualizando pré-visualização da página {page_num}")

            # Renderiza a página selecionada
            doc = fitz.open(pdf_path)
            page = doc.load_page(page_num - 1)
            mat = fitz.Matrix(zoom, zoom)
            pix = page.get_pixmap(matrix=mat)
            doc.close()

            # Converte para imagem PIL
            pil_img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)

            # Ajusta imagem ao tamanho do canvas
            cw = preview_canvas.winfo_width()
            ch = preview_canvas.winfo_height()
            if cw <= 1 or ch <= 1:
                cw, ch = 480, 650
            iw, ih = pil_img.size
            scale = min(cw / iw, ch / ih)
            pil_img = pil_img.resize((int(iw * scale), int(ih * scale)), Image.LANCZOS)

            tk_img = ImageTk.PhotoImage(pil_img)

            # Atualiza o canvas no thread principal
            def update_canvas():
                preview_canvas.delete("all")
                x_center, y_center = cw // 2, ch // 2
                preview_canvas.create_image(x_center, y_center, image=tk_img, anchor="center")
                preview_canvas.image = tk_img
                preview_canvas.full_image = pil_img
                preview_state["current_page"] = page_num
                print(f"[debug] -> Exibindo página {page_num}")

            preview_canvas.after(0, update_canvas)

        except Exception as e:
            print(f"[preview err] {e}")

    # Executa renderização em thread separada (para não travar a interface)
    threading.Thread(target=_render, daemon=True).start()

    
#========================== Marca D'água +===========================================================#

# 1) criação de watermark de texto (força str nos paths)
def make_text_watermark_pdf(text: str, output_path: str, opacity: float = 0.3):
    """
    Cria um PDF de marca d'água com texto vetorial estático (não editável).
    """
    from reportlab.pdfgen import canvas
    from reportlab.lib.colors import Color
    from reportlab.lib.pagesizes import letter

    c = canvas.Canvas(str(output_path), pagesize=letter)
    width, height = letter

    # Define cor azul escuro com opacidade
    dark_blue = Color(0, 0, 0.55, alpha=opacity)

    font_size = 80
    c.setFont("Helvetica-Bold", font_size)
    c.setFillColor(dark_blue)

    # Centraliza o texto
    text_width = c.stringWidth(text, "Helvetica-Bold", font_size)
    x = (width - text_width) / 2
    y = height / 2

    # Rotaciona o texto levemente para parecer marca d'água
    c.saveState()
    c.translate(x, y)
    c.rotate(30)
    c.drawString(0, 0, text)
    c.restoreState()

    c.save()

# 2) criação de watermark por imagem (força str nos paths)
def make_image_watermark_pdf(image_path: str, output_path: str, opacity: float = 0.3, scale: float = 0.8):
    """Cria um PDF de marca d’água a partir de uma imagem PNG."""
    img = Image.open(image_path).convert("RGBA")
    alpha = img.getchannel("A")
    alpha = alpha.point(lambda p: p * opacity)
    img.putalpha(alpha)

    img_stream = io.BytesIO()
    img.save(img_stream, format="PNG")
    img_stream.seek(0)

    c = canvas.Canvas(str(output_path), pagesize=letter)
    width, height = letter
    img_reader = ImageReader(img_stream)

    iw, ih = img.size
    new_w = width * scale
    new_h = ih * (new_w / iw)
    x = (width - new_w) / 2
    y = (height - new_h) / 2
    c.drawImage(img_reader, x, y, width=new_w, height=new_h, mask='auto')
    c.save()


# -------------------- overlay (aplicar watermark PDF sobre cada página) --------------------
def overlay_watermark(pdf_path: str, watermark_pdf: str, out_path: str):
    """
    Sobrepõe `watermark_pdf` (1 página) em cada página de `pdf_path` e grava em out_path.
    Usa abertura em binário para evitar problemas com Path objects.
    """
    pdf_path = str(pdf_path)
    watermark_pdf = str(watermark_pdf)
    out_path = str(out_path)

    # Abre em binário para o PdfReader consumir file-like
    with open(pdf_path, "rb") as f_pdf, open(watermark_pdf, "rb") as f_wm:
        reader = PdfReader(f_pdf)
        watermark = PdfReader(f_wm)
        watermark_page = watermark.pages[0]
        writer = PdfWriter()

        for p in reader.pages:
            # merge_page aplica a watermark (p independente)
            p.merge_page(watermark_page)
            writer.add_page(p)

        with open(out_path, "wb") as fout:
            writer.write(fout)

# 4) wrapper robusto apply_watermark que aceita text / image / watermark_pdf
def apply_watermark(
    pdf_path,
    watermark_image=None,
    watermark_text=None,
    output_path=None,
    scale_factor=1.0,
    alpha=0.5
):
    """Aplica imagem e/ou texto como marca d’água em um PDF."""
    if not output_path:
        raise ValueError("É necessário informar um caminho de saída (output_path).")

    # cria PDF temporário com imagem/texto
    tmp_pdf = Path(tempfile.gettempdir()) / f"__tmp_wm_{os.getpid()}.pdf"
    c = canvas.Canvas(str(tmp_pdf), pagesize=letter)
    width, height = letter

    # 🔹 imagem
    if watermark_image:
        img = Image.open(watermark_image).convert("RGBA")
        new_w = int(img.width * scale_factor)
        new_h = int(img.height * scale_factor)
        img = img.resize((new_w, new_h), resample=Image.LANCZOS)

        # aplica transparência
        if img.mode != "RGBA":
            img = img.convert("RGBA")
        alpha_layer = img.split()[3].point(lambda p: int(p * alpha))
        img.putalpha(alpha_layer)

        buf = io.BytesIO()
        img.save(buf, format="PNG")
        buf.seek(0)
        img_reader = ImageReader(buf)
        x = (width - new_w) / 2
        y = (height - new_h) / 2
        c.drawImage(img_reader, x, y, width=new_w, height=new_h, mask="auto")

    # 🔹 texto
    if watermark_text:
        fs = int(96 * scale_factor)
        c.saveState()
        c.setFont("Helvetica-Bold", fs)
        c.setFillColor(colors.darkblue)
        c.setFillAlpha(alpha)
        text_w = c.stringWidth(watermark_text, "Helvetica-Bold", fs)
        c.drawString((width - text_w) / 2, height / 2, watermark_text)
        c.restoreState()

    c.save()

    # 🔹 aplica marca d’água em cada página
    reader = PdfReader(pdf_path)
    watermark_reader = PdfReader(str(tmp_pdf))
    wm_page = watermark_reader.pages[0]
    writer = PdfWriter()

    for page in reader.pages:
        page.merge_page(wm_page)
        writer.add_page(page)

    with open(output_path, "wb") as f:
        writer.write(f)

    tmp_pdf.unlink(missing_ok=True)
            
# -------------------- geração do PDF de watermark (imagem + texto) --------------------
def make_combined_watermark_pdf(
    image_path: str | None,
    text: str | None,
    out_path: str,
    page_size=letter,
    scale: float = 1.0,
    opacity_img: float = 0.5,
    opacity_text: float = 0.5,
    font_size_base: int = 96,
    text_color=colors.darkblue
):
    """
    Gera um PDF de 1 página com a imagem (redimensionada por `scale` e com `opacity_img`)
    e o texto (com `opacity_text`, cor e tamanho relativo).
    """
    out_path = str(out_path)
    width, height = page_size

    c = canvas.Canvas(out_path, pagesize=page_size)

    # Desenha imagem (se houver)
    if image_path:
        img = Image.open(str(image_path)).convert("RGBA")

        # aplicar escala
        new_w = max(1, int(img.width * scale))
        new_h = max(1, int(img.height * scale))
        img = img.resize((new_w, new_h), resample=Image.LANCZOS)

        # aplicar opacidade na imagem
        if img.mode != "RGBA":
            img = img.convert("RGBA")
        alpha = img.split()[3].point(lambda p: int(p * opacity_img))
        img.putalpha(alpha)

        # salvar em buffer e desenhar
        img_buf = io.BytesIO()
        img.save(img_buf, format="PNG")
        img_buf.seek(0)
        img_reader = ImageReader(img_buf)

        x = (width - new_w) / 2
        y = (height - new_h) / 2
        c.drawImage(img_reader, x, y, width=new_w, height=new_h, mask='auto')

    # Desenha texto por cima (se houver)
    if text:
        fs = int(font_size_base * scale) if scale > 0 else font_size_base
        fs = max(10, fs)
        c.saveState()
        c.setFont("Helvetica-Bold", fs)
        # setFillColor aceita Color; usar RGB com alpha através de setFillAlpha
        try:
            c.setFillColor(text_color)
        except Exception:
            c.setFillColorRGB(0, 0, 0.55)
        c.setFillAlpha(opacity_text)

        text_width = c.stringWidth(str(text), "Helvetica-Bold", fs)
        tx = (width - text_width) / 2
        ty = height / 2
        c.drawString(tx, ty, str(text))
        c.restoreState()

    c.save()

def merge_pdfs(files, out_path):
    writer = PdfWriter()
    for nome_arquivo in files:
        reader = PdfReader(nome_arquivo)
        for pagina in reader.pages:
            writer.add_page(pagina)
    with open(out_path, "wb") as f:
        writer.write(f)


def reorder_pdf(input_path: str, page_order: List[int], out_path: str):
    reader = PdfReader(input_path)
    writer = PdfWriter()
    total = len(reader.pages)
    # page_order are 1-based indices; negative or out-of-range will be ignored
    for idx in page_order:
        if 1 <= idx <= total:
            writer.add_page(reader.pages[idx - 1])
    with open(out_path, "wb") as f:
        writer.write(f)


def split_pdf_ranges(input_path: str, ranges: List[Tuple[int, int]], out_dir: str):
    reader = PdfReader(input_path)
    total = len(reader.pages)
    basename = Path(input_path).stem
    outputs = []
    for i, (a, b) in enumerate(ranges, start=1):
        a0 = max(1, a)
        b0 = min(total, b)
        if a0 > b0:
            continue
        writer = PdfWriter()
        for p in range(a0 - 1, b0):
            writer.add_page(reader.pages[p])
        out_path = Path(out_dir) / f"{basename}_part{i}.pdf"
        with open(out_path, "wb") as f:
            writer.write(f)
        outputs.append(str(out_path))
    return outputs


def images_to_pdfs(image_paths: List[str], out_dir: str, single_pdf=False, single_name="images_merged.pdf") -> List[str]:
    """Converte cada imagem em um PDF (ou junta todas em um único PDF).
    Retorna lista de arquivos gerados."""
    outputs = []
    pil_images = []
    for img_path in image_paths:
        img = Image.open(img_path)
        # Convert to RGB if necessary
        if img.mode in ("RGBA", "LA"):
            bg = Image.new("RGB", img.size, (255, 255, 255))
            bg.paste(img, mask=img.split()[3])
            img = bg
        else:
            img = img.convert("RGB")

        if single_pdf:
            pil_images.append(img)
        else:
            name = Path(img_path).stem + ".pdf"
            out_path = Path(out_dir) / name
            img.save(out_path, "PDF", resolution=100.0)
            outputs.append(str(out_path))

    if single_pdf and pil_images:
        first, rest = pil_images[0], pil_images[1:]
        out_path = Path(out_dir) / single_name
        first.save(out_path, "PDF", save_all=True, append_images=rest)
        outputs.append(str(out_path))

    return outputs


def docx_to_pdf_batch(docx_paths: List[str], out_dir: str) -> List[str]:
    outputs = []
    # If docx2pdf available, use it (Windows: MS Word backend).
    if docx2pdf_convert is not None:
        for p in docx_paths:
            try:
                basename = Path(p).stem + ".pdf"
                out_path = Path(out_dir) / basename
                # docx2pdf convert(in, out) where out can be file path
                docx2pdf_convert(p, str(out_path))
                outputs.append(str(out_path))
            except Exception as e:
                print("docx2pdf failed for", p, e)
    else:
        # Try LibreOffice soffice (cross-platform) via subprocess
        soffice = shutil.which("soffice") or shutil.which("libreoffice")
        if soffice:
            for p in docx_paths:
                try:
                    # soffice --headless --convert-to pdf --outdir OUTDIR INPUT
                    subprocess.run([soffice, "--headless", "--convert-to", "pdf", "--outdir", str(out_dir), str(p)], check=True)
                    outputs.append(str(Path(out_dir) / (Path(p).stem + ".pdf")))
                except Exception as e:
                    print("soffice failed for", p, e)
        else:
            raise RuntimeError("Nenhum conversor .docx->pdf disponível (docx2pdf ausente e soffice não encontrado).")
    return outputs

def mix_media_to_pdf(self):
    import fitz  # PyMuPDF
    from PIL import Image

    filepaths = filedialog.askopenfilenames(
        title="Selecione imagens e/ou PDFs",
        filetypes=[
            ("Imagens e PDFs", "*.jpg *.jpeg *.png *.bmp *.tif *.tiff *.pdf"),
            ("Todos os arquivos", "*.*"),
        ],
    )
    if not filepaths:
        return

    output_path = filedialog.asksaveasfilename(
        defaultextension=".pdf",
        filetypes=[("Arquivo PDF", "*.pdf")],
        title="Salvar PDF combinado como...",
    )
    if not output_path:
        return

    pdf = fitz.open()

    for fp in filepaths:
        try:
            ext = Path(fp).suffix.lower()
            if ext == ".pdf":
                with fitz.open(fp) as src_pdf:
                    for page in src_pdf:
                        pdf.insert_pdf(src_pdf, from_page=page.number, to_page=page.number)
            else:
                # Converter imagem em PDF temporariamente
                img = Image.open(fp).convert("RGB")
                img_bytes = img.tobytes()
                rect = fitz.Rect(0, 0, img.width, img.height)
                img_page = pdf.new_page(width=img.width, height=img.height)
                pix = fitz.Pixmap(fitz.csRGB, img.width, img.height, img_bytes)
                img_page.insert_image(rect, pixmap=pix)
                pix = None
        except Exception as e:
            print(f"Erro ao processar {fp}: {e}")

    pdf.save(output_path)
    pdf.close()

    messagebox.showinfo("Sucesso", f"PDF combinado salvo em:\n{output_path}")

#======== Auxiliar Dividir PDF =============================================================================================
def save_organized_pdf(self, win):
    import fitz
    from tkinter import filedialog, messagebox

    path = filedialog.asksaveasfilename(defaultextension=".pdf")
    if not path:
        return

    with fitz.open(self._organize_pdf_path) as doc:
        new = fitz.open()
        for i in self._orig_page_idx:
            new.insert_pdf(doc, from_page=i, to_page=i)
        new.save(path)
        new.close()

    messagebox.showinfo("Sucesso", "PDF salvo.")
    win.destroy()

def execute_split_visual(self, mode, interval_text, every_text):
    # No próximo passo vamos ligar isso à ordem visual
    print("Modo:", mode)

#======== Auxiliar PDF para Imagem =============================================================================================
def convert_pdfs_to_images(self, pdf_files, output_dir, img_format):
    def worker():
        for pdf_path in pdf_files:
            pdf_name = os.path.splitext(os.path.basename(pdf_path))[0]

            with fitz.open(pdf_path) as doc:
                for i, page in enumerate(doc):
                    pix = page.get_pixmap(dpi=200)

                    ext = "png" if img_format == "png" else "jpg"
                    output_path = os.path.join(
                        output_dir,
                        f"{pdf_name}_P{i+1}.{ext}"
                    )

                    if img_format == "jpg":
                        pix.save(output_path, output="jpeg")
                    else:
                        pix.save(output_path)

        self.after(0, lambda: messagebox.showinfo(
            "Concluído",
            "Conversão finalizada com sucesso."
        ))

    threading.Thread(target=worker, daemon=True).start()
    

###################### INTERFACE DE USUÁRIO E FUNÇÕES DAS TELAS DO APP #################################################

class App(ctk.CTk):
    def __init__(self):
        super().__init__()

        # ===== Configuração da janela principal =====
        self.title(APP_NAME)
        self.geometry("1080x640")  # tamanho mais justo e equilibrado
        self.minsize(960, 580)
        ctk.set_appearance_mode("dark")

        # ===== Frame principal com 2 colunas: lateral + conteúdo =====
        self.grid_columnconfigure(0, weight=0)  # barra lateral
        self.grid_columnconfigure(1, weight=1)  # área principal
        self.grid_rowconfigure(0, weight=1)

        # ===== Barra lateral =====
        self.sidebar = ctk.CTkFrame(self, width=200, corner_radius=0)
        self.sidebar.grid(row=0, column=0, sticky="nswe")
        self.sidebar.grid_rowconfigure(8, weight=1)  # empurra o rodapé

        ctk.CTkLabel(
            self.sidebar, text=APP_NAME, font=ctk.CTkFont(size=18, weight="bold")
        ).grid(row=0, column=0, pady=(20, 10), padx=20)

        # Botões da barra lateral
        btn_cfg = dict(width=180, height=40, font=ctk.CTkFont(size=13, weight="bold"))

        # Cria um frame interno para agrupar os botões e manter o espaçamento consistente
        btn_frame = ctk.CTkFrame(self.sidebar, fg_color="transparent")
        btn_frame.grid(row=1, column=0, pady=(20, 20), padx=10, sticky="n")

        buttons = [
            ("🏠 Início", self.open_home),
            ("Marca d'água", self.open_tab_watermark),
            ("Juntar PDFs", self.open_tab_merge),
            ("Organizar PDFs", self.open_tab_organize),
            ("Dividir PDF", self.open_tab_split_pdf),
            ("Imagens → PDF", self.open_tab_images),
            ("PDF → Imagens", self.open_pdf_to_images),
            ("Word → PDF", self.open_tab_word),
            ("PDF → Word", self.open_tab_pdf_to_word),
            ("Juntar Imagens e PDFs", self.open_tab_mix_media),
            ("Combinar e Organizar", self.start_organize_ui_custom),
            ("Renomeador de PDFs", self.open_tab_renamer),
        ]

        for i, (text, cmd) in enumerate(buttons):
            ctk.CTkButton(btn_frame, text=text, command=cmd, **btn_cfg).grid(
                row=i, column=0, pady=5
            )


        # Rodapé
        ctk.CTkLabel(
            self.sidebar,
            text="© Concremat Engenharia\nDesenvolvido por Gustavo Dumont",
            font=ctk.CTkFont(size=10),
            justify="center",
            text_color="#888",
        ).grid(row=9, column=0, pady=(40, 10), padx=10, sticky="s")


        # Rodapé
        self.sidebar.grid_rowconfigure(9, weight=1)
        ctk.CTkLabel(
            self.sidebar,
            text="© Concremat Engenharia\nDesenvolvido por Gustavo Dumont",
            font=ctk.CTkFont(size=10),
            justify="center",
            text_color="#888",
        ).grid(row=9, column=0, pady=(40, 10), padx=10, sticky="s")


        # ===== Área principal =====
        self.main = ctk.CTkFrame(self, corner_radius=0)
        self.main.grid(row=0, column=1, sticky="nswe", padx=0, pady=0)
        self.current_frame = None

        # ===== Abre a tela inicial =====
        self.open_home()

    # ---------------------------------------------------------
    # Utilitário para limpar a tela atual
    # ---------------------------------------------------------
    def clear_main(self):
        if self.current_frame is not None:
            self.current_frame.destroy()
            self.current_frame = None

    # ---------------------------------------------------------
    # Exibe mensagens de status no rodapé
    # ---------------------------------------------------------
    def set_status(self, message):
        if hasattr(self, "status_label"):
            self.status_label.configure(text=message)
        else:
            print(message)        

    # ---------------------------------------------------------
    # Tela inicial (home)
    # ---------------------------------------------------------
    def open_home(self):
        self.clear_main()
        frame = ctk.CTkFrame(self.main)
        frame.pack(expand=True)
        self.current_frame = frame

        # ===== Logo =====
        if getattr(sys, "frozen", False):
            base_dir = Path(sys.executable).parent
        else:
            base_dir = Path(os.getcwd())

        # Caminho completo da logo
        logo_path = base_dir / "Logo.png"
        logo_img_tk = None

        # ===== Tenta carregar a imagem ========
        if logo_path.exists():
            try:
                img = Image.open(logo_path).convert("RGBA")
                img = img.resize((194, 194), Image.LANCZOS)
                logo_img_tk = ImageTk.PhotoImage(img)
            except Exception as e:
                print("Erro ao carregar logo:", e)
                logo_img_tk = None
        else:
            print(f"Logo não encontrada em: {logo_path}")

        # ===== Exibe a imagem =================
        if logo_img_tk:
            logo_lbl = ctk.CTkLabel(frame, image=logo_img_tk, text="")
            logo_lbl.image = logo_img_tk
            logo_lbl.pack(pady=(20, 10))
        else:
            ctk.CTkLabel(frame, text="[Logo não encontrada]").pack(pady=10)

        # ===== Título e subtítulo ==============
        ctk.CTkLabel(
            frame, text=APP_NAME, font=ctk.CTkFont(size=26, weight="bold")
        ).pack(pady=(5, 2))

        subtitle = (
            "Gustavo Freitas Gomes Dumont\n"
            "Núcleo Administrativo Indústria e Mineração · Concremat Engenharia\n"
            "Novembro de 2025"
        )
        ctk.CTkLabel(
            frame, text=subtitle, justify="center", font=ctk.CTkFont(size=12)
        ).pack(pady=(5, 15))

        # ===== Botões principais ===============
        btn_cfg = dict(width=300, height=50, font=ctk.CTkFont(size=14, weight="bold"))
        ctk.CTkButton(
            frame, text="💧 Marca d'água", command=self.open_tab_watermark, **btn_cfg
        ).pack(pady=6)
        ctk.CTkButton(
            frame, text="📎 Juntar PDFs", command=self.open_tab_merge, **btn_cfg
        ).pack(pady=6)
        ctk.CTkButton(
            frame, text="🗂️ Organizar PDFs", command=self.open_tab_organize, **btn_cfg
        ).pack(pady=6)
        ctk.CTkButton(
            frame, text="🖼️ Imagens → PDF", command=self.open_tab_images, **btn_cfg
        ).pack(pady=6)
        ctk.CTkButton(
            frame, text="📄 Word → PDF", command=self.open_tab_word, **btn_cfg
        ).pack(pady=6)

        ctk.CTkLabel(
            frame, text=f"{APP_NAME}", text_color="#888", font=ctk.CTkFont(size=10)
        ).pack(pady=(20, 5))        

    # ================= Tab Marca D'água ======================================================================
    def open_tab_watermark(self):
        self.clear_main()
        frame = ctk.CTkFrame(self.main)
        frame.grid(row=0, column=0, sticky="nswe")
        self.current_frame = frame

        # Título
        lbl = ctk.CTkLabel(
            frame,
            text="Aplicar Marca d'água em Lote",
            font=ctk.CTkFont(size=20, weight="bold")
        )
        lbl.grid(row=0, column=0, padx=20, pady=(20, 10), sticky="w")

        # Instruções
        instructions = (
            "📄 Instruções de uso:\n\n"
            "1️⃣ Selecione os arquivos PDF que deseja processar.\n"
            "2️⃣ Escolha uma imagem (PNG/JPG) e/ou insira um texto como marca d'água.\n"
            "3️⃣ Ajuste o tamanho e a transparência da imagem conforme desejar.\n"
            "4️⃣ Clique no botão abaixo para aplicar em todos os PDFs.\n\n"
            "💾 Os arquivos resultantes serão salvos na mesma pasta original, "
            "com o sufixo '_watermarked'."
        )

        lbl_instr = ctk.CTkLabel(
            frame,
            text=instructions,
            justify="left",
            wraplength=580,
            font=ctk.CTkFont(size=14)
        )
        lbl_instr.grid(row=1, column=0, padx=20, pady=(0, 20), sticky="w")

        # Controle de tamanho
        ctk.CTkLabel(
            frame,
            text="Tamanho da imagem da marca d'água:",
            font=ctk.CTkFont(size=14, weight="bold")
        ).grid(row=2, column=0, padx=20, pady=(0, 6), sticky="w")

        self.image_scale_var = ctk.DoubleVar(value=1.0)

        scale_frame = ctk.CTkFrame(frame, fg_color="transparent")
        scale_frame.grid(row=3, column=0, padx=20, pady=(0, 20), sticky="we")

        self.scale_label = ctk.CTkLabel(scale_frame, text="100%", width=60)
        self.scale_label.pack(side="right", padx=6)

        scale = ctk.CTkSlider(
            scale_frame,
            from_=0.01,
            to=3.0,
            number_of_steps=299,
            variable=self.image_scale_var,
            command=lambda val: self.scale_label.configure(text=f"{int(float(val)*100)}%")
        )
        scale.pack(fill="x", expand=True, padx=(0, 6))

        # Controle de transparência
        ctk.CTkLabel(
            frame,
            text="Transparência da marca d'água:",
            font=ctk.CTkFont(size=14, weight="bold")
        ).grid(row=4, column=0, padx=20, pady=(0, 6), sticky="w")

        self.image_alpha_var = ctk.DoubleVar(value=0.5)

        alpha_frame = ctk.CTkFrame(frame, fg_color="transparent")
        alpha_frame.grid(row=5, column=0, padx=20, pady=(0, 20), sticky="we")

        self.alpha_label = ctk.CTkLabel(alpha_frame, text="50%", width=60)
        self.alpha_label.pack(side="right", padx=6)

        alpha_slider = ctk.CTkSlider(
            alpha_frame,
            from_=0.0,
            to=1.0,
            number_of_steps=100,
            variable=self.image_alpha_var,
            command=lambda val: self.alpha_label.configure(text=f"{int(float(val)*100)}%")
        )
        alpha_slider.pack(fill="x", expand=True, padx=(0, 6))

        # Botão principal
        btn_run = ctk.CTkButton(
            frame,
            text="🚀 Aplicar Marca d'água em Massa",
            font=ctk.CTkFont(size=16, weight="bold"),
            height=48,
            width=300,
            fg_color="#1f6aa5",
            hover_color="#15527a",
            command=self.start_watermark
        )
        btn_run.grid(row=6, column=0, padx=20, pady=(10, 20), sticky="we")


    def select_watermark_files(self):
        paths = filedialog.askopenfilenames(title="Selecionar PDFs", filetypes=[("PDF files", "*.pdf")])
        if paths:
            self.wm_files = list(paths)
            self.wm_files_var.set(f"{len(self.wm_files)} arquivos selecionados")

    def select_watermark_pdf(self):
        path = filedialog.askopenfilename(title="Selo PDF (opcional)", filetypes=[("PDF files","*.pdf")])
        if path:
            self.wm_watermark_pdf = path
            self.wm_pdf_var.set(Path(path).name)

    def select_watermark_image(self):
        path = filedialog.askopenfilename(
            title="Imagem da marca d'água (PNG ou JPG)",
            filetypes=[("Imagens", "*.png;*.jpg;*.jpeg")]
        )
        if path:
            self.wm_watermark_img = path
            self.wm_img_var.set(Path(path).name)
        

    def select_wm_outdir(self):
        d = filedialog.askdirectory(title="Pasta de saída")
        if d:
            self.wm_outdir.set(d)

    def start_watermark(self):
        try:
            files = filedialog.askopenfilenames(filetypes=[("Arquivos PDF", "*.pdf")])
            if not files:
                return

            out_dir = filedialog.askdirectory(title="Selecione a pasta de destino")
            if not out_dir:
                return

            wm_input = filedialog.askopenfilename(
                title="Selecione a imagem da marca d'água (opcional)",
                filetypes=[("Imagem PNG", "*.png"), ("Todos os arquivos", "*.*")]
            )

            wm_text = simpledialog.askstring(
                "Texto da marca d'água",
                "Se desejar, escreva algo para complementar a marca d'água:"
            )

            for f in files:
                output_path = Path(out_dir) / f"WM_{Path(f).name}"
                self.set_status(f"Aplicando marca d'água em {Path(f).name}...")

                try:
                    apply_watermark(
                        pdf_path=f,
                        watermark_image=wm_input if wm_input else None,
                        watermark_text=wm_text if wm_text else None,
                        output_path=output_path,
                        scale_factor=self.image_scale_var.get(),
                        alpha=self.image_alpha_var.get()
                    )
                except Exception as e:
                    messagebox.showerror("Erro", f"Falha ao aplicar marca d'água em {f}:\n{e}")
                    continue

            self.set_status("Marca d'água aplicada com sucesso.")
            messagebox.showinfo("Sucesso", "Marca d'água aplicada com sucesso a todos os PDFs.")

        except Exception as e:
            messagebox.showerror("Erro", f"Falha geral ao aplicar marca d'água:\n{e}")



    def _do_watermark(self, files, outdir_str, text, image_path, watermark_pdf):
        """
        Aplica watermark a cada arquivo da lista `files`.
        Recebe outdir como string para evitar WindowsPath vazando.
        """
        self.set_status("Aplicando watermark...")
        try:
            outdir = Path(outdir_str)
        except Exception:
            outdir = Path(str(outdir_str))

        self.wm_progress.set(0)
        n = len(files)
        for i, f in enumerate(files, start=1):
            try:
                # garante strings para a função
                infile = str(f)
                out_file = outdir / (Path(f).stem + "_wm.pdf")
                out_file_str = str(out_file)

                # chama apply_watermark com keyword args para não confundir ordem
                apply_watermark(
                    input_pdf=infile,
                    output_pdf=out_file_str,
                    text=text if text else None,
                    image_path=str(image_path) if image_path else None,
                    watermark_pdf=str(watermark_pdf) if watermark_pdf else None,
                    opacity=0.12
                )
            except Exception as e:
                # loga o erro e segue (não interrompe o lote)
                print("Erro watermark em", f, e)
            self.wm_progress.set(i / n)

        self.wm_progress.set(1.0)
        self.set_status("Concluído: watermark aplicado")


    # ================== Unir PDFs Tab ============================================
    def open_tab_merge(self):
        self.clear_main()
        frame = ctk.CTkFrame(self.main)
        frame.grid(row=0, column=0, sticky="nswe")
        self.current_frame = frame

        # --- Título ---
        lbl = ctk.CTkLabel(
            frame,
            text="Juntar PDFs em Lote",
            font=ctk.CTkFont(size=20, weight="bold")
        )
        lbl.grid(row=0, column=0, padx=20, pady=(20, 10), sticky="w")

        # --- Instruções simplificadas ---
        instructions = (
            "📄 Instruções de uso:\n\n"
            "1️⃣ Selecione os arquivos PDF que deseja juntar.\n"
            "2️⃣ A ordem seguirá a seleção feita no explorador de arquivos.\n"
            "3️⃣ Clique no botão abaixo para gerar um único PDF combinado.\n\n"
            "💾 O arquivo final será salvo na pasta do primeiro PDF selecionado "
            "com o nome 'PDFs_Unidos.pdf'."
        )

        lbl_instr = ctk.CTkLabel(
            frame,
            text=instructions,
            justify="left",
            wraplength=580,
            font=ctk.CTkFont(size=14)
        )
        lbl_instr.grid(row=1, column=0, padx=20, pady=(0, 20), sticky="w")

        # --- Barra de progresso (adicionar em open_tab_merge) ---
        self.merge_progress = ctk.CTkProgressBar(frame)
        self.merge_progress.grid(row=2, column=0, columnspan=1, sticky="ew", padx=20, pady=(8, 12))
        # inicializa em zero
        self.merge_progress.set(0.0)

        # --- Botão principal ---
        btn_run = ctk.CTkButton(
            frame,
            text="🧩 Juntar PDFs",
            font=ctk.CTkFont(size=16, weight="bold"),
            height=48,
            width=300,
            fg_color="#1f6aa5",
            hover_color="#15527a",
            command=self.start_merge
        )
        btn_run.grid(row=2, column=0, padx=20, pady=(10, 20), sticky="we")


    def select_merge_files(self):
        paths = filedialog.askopenfilenames(title="Selecionar PDFs para juntar", filetypes=[("PDF", "*.pdf")])
        if paths:
            self.merge_files = list(paths)
            self.merge_files_var.set(f"{len(self.merge_files)} arquivos selecionados")

    def select_merge_outdir(self):
        d = filedialog.askdirectory(title="Pasta de saída")
        if d:
            self.merge_outdir.set(d)

    def start_merge(self):
        # 🔹 Passo 1: Selecionar arquivos
        files = filedialog.askopenfilenames(
            title="Selecione os arquivos PDF para juntar",
            filetypes=[("Arquivos PDF", "*.pdf")]
        )
        if not files:
            messagebox.showwarning("Aviso", "Nenhum PDF selecionado.")
            return

        # 🔹 Passo 2: Definir pasta de saída automaticamente
        first_pdf = Path(files[0])
        outdir = first_pdf.parent
        outdir.mkdir(parents=True, exist_ok=True)

        # 🔹 Passo 3: Nome do arquivo final
        out_path = outdir / "PDFs_Unidos.pdf"

        # 🔹 Passo 4: Thread para processamento (não trava interface)
        t = threading.Thread(target=self._do_merge, args=(files, str(out_path)), daemon=True)
        t.start()

    def _do_merge(self, files, out_path):
        self.set_status("Juntando PDFs...")
        try:
            merge_pdfs(files, out_path)
            self.set_status(f"Concluído: {out_path}")
            self.merge_progress.set(1.0)
        except Exception as e:
            messagebox.showerror("Erro", str(e))
            self.set_status("Erro ao juntar PDFs")


    # ============================= Organizar PDFs Tab ==============================
    def open_tab_organize(self):
        self.clear_main()
        frame = ctk.CTkFrame(self.main)
        frame.grid(row=0, column=0, sticky="nswe")
        self.current_frame = frame

        lbl = ctk.CTkLabel(frame, text="Organizar Páginas de PDFs", font=ctk.CTkFont(size=18, weight="bold"))
        lbl.pack(pady=(15, 10))

        instructions = (
            "Instruções de uso:\n"
            "1. Clique no botão abaixo para selecionar o PDF que deseja organizar.\n"
            "2. Uma nova janela será aberta mostrando todas as páginas do arquivo.\n"
            "3. Clique em uma miniatura para visualizar à direita.\n"
            "4. Arraste e solte as páginas na ordem desejada.\n"
            "5. Clique em 'Salvar' para gerar o novo PDF organizado."
        )
        ctk.CTkLabel(frame, text=instructions, justify="left", wraplength=600).pack(padx=20, pady=(0, 20))

        ctk.CTkButton(
            frame,
            text="Organizar Páginas de PDFs",
            height=50,
            font=ctk.CTkFont(size=15, weight="bold"),
            command=self.start_organize_ui
        ).pack(pady=20)


    def start_organize_ui(self):
        import threading
        from pathlib import Path
        from PyPDF2 import PdfReader, PdfWriter
        from PIL import Image, ImageTk, ImageDraw, ImageFont
        import tkinter as tk
        from tkinter import filedialog, messagebox
        import customtkinter as ctk
        import fitz  # PyMuPDF

        # --- Seleção de arquivo ---
        file_path = filedialog.askopenfilename(
            title="Selecionar PDF para organizar",
            filetypes=[("Arquivos PDF", "*.pdf")]
        )
        if not file_path:
            return

        # --- Janela ---
        org_window = ctk.CTkToplevel(self)
        org_window.title(f"Organizar páginas - {Path(file_path).name}")
        org_window.geometry("1280x800")

        frame_main = ctk.CTkFrame(org_window)
        frame_main.pack(fill="both", expand=True, padx=10, pady=10)

        # Layout: left = miniaturas (scroll), right = preview + zoom
        left_frame = ctk.CTkFrame(frame_main)
        left_frame.pack(side="left", fill="both", expand=True, padx=(0, 10))

        right_frame = ctk.CTkFrame(frame_main, width=520)
        right_frame.pack(side="right", fill="y")

        ctk.CTkLabel(right_frame, text="Pré-visualização", font=ctk.CTkFont(size=16, weight="bold")).pack(pady=10)
        preview_canvas = tk.Canvas(right_frame, bg="#1e1e1e", width=480, height=650, highlightthickness=0)
        preview_canvas.pack(padx=10, pady=(0, 10))
        preview_canvas.bind("<Configure>", lambda e: preview_canvas.update_idletasks())

        # estado local para preview (controla cancelamento + thread atual)
        preview_state = {"current_page": None, "render_thread": None, "cancel_flag": False}

        # estado de arrasto (para distinguir click vs drag)
        drag_state = {"start_x": 0, "start_y": 0, "widget": None, "index": None, "moved": False}

        # lista de widgets miniatura (closure)
        thumb_widgets = []

        # --- Helpers: rasterização com fitz (retorna PIL.Image) ---
        def convert_page_silent(pdf_path_local, page_number, target_width=None, dpi=150):
            try:
                with fitz.open(pdf_path_local) as doc:
                    page = doc.load_page(page_number - 1)
                    mat = fitz.Matrix(dpi / 72.0, dpi / 72.0)
                    pix = page.get_pixmap(matrix=mat, alpha=False)
                    img = Image.frombytes("RGB", (pix.width, pix.height), pix.samples)
                if target_width:
                    w, h = img.size
                    new_h = int(h * (target_width / w))
                    img = img.resize((int(target_width), new_h), Image.LANCZOS)
                return img
            except Exception as e:
                print(f"[convert_page_silent err] pág {page_number}: {e}")
                return None

        # overlay com número da página
        def add_page_number_overlay(pil_img, page_number):
            draw = ImageDraw.Draw(pil_img)
            font_size = max(14, int(pil_img.width * 0.06))
            try:
                font = ImageFont.truetype("arial.ttf", font_size)
            except Exception:
                font = ImageFont.load_default()
            text = str(page_number)
            try:
                bbox = draw.textbbox((0, 0), text, font=font)
                tw, th = bbox[2] - bbox[0], bbox[3] - bbox[1]
            except Exception:
                tw, th = draw.textsize(text, font=font)
            draw.rectangle([(0, 0), (tw + 12, th + 12)], fill=(0, 0, 0, 150))
            draw.text((6, 6), text, fill="white", font=font)
            return pil_img

        # --- safe_update_preview: garante que só renderiza a página clicada ---
        def safe_update_preview(pdf_path_local, page, preview_canvas_obj, zoom=1.5):
            # cancela renderização anterior
            if preview_state.get("render_thread") and preview_state["render_thread"].is_alive():
                preview_state["cancel_flag"] = True

            # prepara novo estado
            preview_state["cancel_flag"] = False
            preview_state["current_page"] = page

            # feedback imediato
            preview_canvas_obj.delete("all")
            preview_canvas_obj.create_text(
                preview_canvas_obj.winfo_reqwidth() // 2,
                preview_canvas_obj.winfo_reqheight() // 2,
                text=f"Carregando página {page}...",
                fill="gray",
                font=("Arial", 12, "italic")
            )
            preview_canvas_obj.update_idletasks()

            def _render():
                try:
                    doc = fitz.open(pdf_path_local)
                    if preview_state["cancel_flag"]:
                        doc.close()
                        return

                    page_obj = doc.load_page(page - 1)
                    mat = fitz.Matrix(zoom, zoom)
                    pix = page_obj.get_pixmap(matrix=mat)
                    doc.close()

                    if preview_state["cancel_flag"]:
                        return

                    pil_img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)

                    # redimensiona para caber no canvas mantendo proporção
                    cw = preview_canvas_obj.winfo_width()
                    ch = preview_canvas_obj.winfo_height()
                    if cw <= 1 or ch <= 1:
                        cw, ch = 480, 650
                    iw, ih = pil_img.size
                    scale = min(cw / iw, ch / ih)
                    pil_disp = pil_img.resize((max(1, int(iw * scale)), max(1, int(ih * scale))), Image.LANCZOS)
                    tk_img = ImageTk.PhotoImage(pil_disp)

                    def update_canvas():
                        if preview_state["cancel_flag"]:
                            return
                        preview_canvas_obj.delete("all")
                        x_center, y_center = cw // 2, ch // 2
                        preview_canvas_obj.create_image(x_center, y_center, image=tk_img, anchor="center")
                        # manter referência pra evitar GC
                        preview_canvas_obj.image = tk_img
                        preview_canvas_obj.full_image = pil_img
                        preview_state["current_page"] = page
                        # debug:
                        # print(f"[debug] -> Exibindo página {page}")

                    preview_canvas_obj.after(0, update_canvas)

                except Exception as e:
                    print(f"[preview err] {e}")

            t = threading.Thread(target=_render, daemon=True)
            preview_state["render_thread"] = t
            t.start()

        # --- Funções de clique / drag nas miniaturas ---
        def on_thumb_press(event, page):
            widget = event.widget
            try:
                idx = thumb_widgets.index(widget)
            except ValueError:
                idx = None
            drag_state.update({
                "start_x": event.x_root,
                "start_y": event.y_root,
                "widget": widget,
                "index": idx,
                "moved": False
            })
            widget.config(relief="ridge")

        def on_thumb_motion(event):
            if not drag_state["widget"]:
                return
            dx = abs(event.x_root - drag_state["start_x"])
            dy = abs(event.y_root - drag_state["start_y"])
            if dx > 8 or dy > 8:
                drag_state["moved"] = True
                # opcional: você pode mover uma cópia visual aqui (não implementado) 

        def on_thumb_release(event):
            widget = drag_state["widget"]
            if not widget:
                return
            widget.config(relief="solid")

            if not drag_state["moved"]:
                page = getattr(widget, "page_num", None)
                if page is not None:
                    # clique simples -> só renderiza a página clicada
                    safe_update_preview(file_path, page, preview_canvas)
            else:
                # reorganiza por posição do release
                try:
                    x = event.x_root - inner_frame.winfo_rootx()
                    y = event.y_root - inner_frame.winfo_rooty()
                    dest_col = max(0, int(x // 280))
                    dest_row = max(0, int(y // 360))
                    dest_index = min(dest_row * 3 + dest_col, len(thumb_widgets) - 1)

                    src_index = drag_state["index"]
                    if src_index is None:
                        drag_state.update({"widget": None, "index": None, "moved": False})
                        return
                    item = thumb_widgets.pop(src_index)
                    thumb_widgets.insert(dest_index, item)
                    for idx, w in enumerate(thumb_widgets):
                        w.grid(row=idx // 3, column=idx % 3, padx=8, pady=8)
                except Exception as e:
                    print(f"[drag release err] {e}")

            drag_state.update({"widget": None, "index": None, "moved": False})

        # --- Cria miniatura (widget) de forma isolada ---
        def create_thumbnail_button(parent, pil_image, page_num):
            try:
                tk_img = ImageTk.PhotoImage(pil_image)
                lbl = tk.Label(
                    parent,
                    image=tk_img,
                    bg="#2b2b2b",
                    bd=2,
                    relief="solid",
                    cursor="hand2"
                )
                lbl.image = tk_img
                lbl.page_num = page_num
                lbl.grid(row=(page_num - 1) // 3, column=(page_num - 1) % 3, padx=8, pady=8)

                # bind separados: press, motion e release para distinguir drag/click
                lbl.bind("<ButtonPress-1>", lambda e, p=page_num: on_thumb_press(e, p))
                lbl.bind("<B1-Motion>", on_thumb_motion)
                lbl.bind("<ButtonRelease-1>", on_thumb_release)

                # hover visual
                lbl.bind("<Enter>", lambda e: lbl.config(bd=3, relief="ridge"))
                lbl.bind("<Leave>", lambda e: lbl.config(bd=2, relief="solid"))
                return lbl
            except Exception as e:
                print(f"[create_thumbnail_button err] {e}")
                return None

        # --- Scroll area para miniaturas ---
        canvas = tk.Canvas(left_frame, bg="#1e1e1e", highlightthickness=0)
        scrollbar = ctk.CTkScrollbar(left_frame, orientation="vertical", command=canvas.yview)
        canvas.configure(yscrollcommand=scrollbar.set)

        scrollbar.pack(side="right", fill="y")
        canvas.pack(side="left", fill="both", expand=True, padx=(0, 8))

        inner_frame = tk.Frame(canvas, bg="#2b2b2b")
        canvas_window = canvas.create_window((0, 0), window=inner_frame, anchor="nw")

        def update_scrollregion(event=None):
            canvas.configure(scrollregion=canvas.bbox("all"))
        inner_frame.bind("<Configure>", update_scrollregion)

        def _on_mousewheel(self, event):
            if hasattr(self, "canvas_frame") and hasattr(self.canvas_frame, "canvas"):
                self.canvas_frame.canvas.yview_scroll(int(-1 * (event.delta / 40)), "units")


        # bind enter/leave para ativar scroll quando o mouse estiver sobre as miniaturas
        canvas.bind("<Enter>", lambda e: org_window.bind_all("<MouseWheel>", _on_mousewheel))
        canvas.bind("<Leave>", lambda e: org_window.unbind_all("<MouseWheel>"))

        # --- Gera miniaturas em thread (não bloqueia UI) ---
        reader = PdfReader(file_path)
        n_pages = len(reader.pages)

        def generate_thumbnails():
            for i in range(n_pages):
                try:
                    page_index = i + 1
                    pil = convert_page_silent(file_path, page_index, target_width=260, dpi=90)
                    if pil is None:
                        continue
                    pil = add_page_number_overlay(pil, page_index)

                    # cria widget no thread principal
                    def mk(i=page_index, pil_local=pil):
                        lbl = create_thumbnail_button(inner_frame, pil_local, i)
                        if lbl:
                            thumb_widgets.append(lbl)
                            update_scrollregion()
                    org_window.after(0, mk)

                except Exception as e:
                    print(f"[thumb err] página {i+1}: {e}")

            # garante scrollregion final
            org_window.after(100, update_scrollregion)

        threading.Thread(target=generate_thumbnails, daemon=True).start()

        # --- Funções de salvar PDF reorganizado ---
        def save_new_pdf():
            try:
                order_page_nums = [w.page_num for w in thumb_widgets]  # 1-based
                out_path = filedialog.asksaveasfilename(
                    defaultextension=".pdf",
                    filetypes=[("Arquivos PDF", "*.pdf")],
                    title="Salvar PDF organizado"
                )
                if not out_path:
                    return
                writer = PdfWriter()
                for idx in order_page_nums:
                    writer.add_page(reader.pages[idx - 1])
                with open(out_path, "wb") as f:
                    writer.write(f)
                messagebox.showinfo("Sucesso", f"PDF salvo em:\n{out_path}")
                org_window.destroy()
            except Exception as e:
                messagebox.showerror("Erro", str(e))

        ctk.CTkButton(org_window, text="Salvar PDF Organizado", command=save_new_pdf).pack(pady=12)

    # ========================= Imagens Tab =======================================================================================#
    def open_tab_images(self):
        self.clear_main()
        frame = ctk.CTkFrame(self.main)
        frame.grid(row=0, column=0, sticky="nswe")
        self.current_frame = frame

        lbl = ctk.CTkLabel(frame, text="Converter imagens em PDF (massa)", font=ctk.CTkFont(size=16, weight="bold"))
        lbl.grid(row=0, column=0, sticky="w", padx=12, pady=8)

        btn_select = ctk.CTkButton(frame, text="Selecionar imagens", command=self.select_image_files)
        btn_select.grid(row=1, column=0, padx=12, pady=6, sticky="w")
        self.img_files_var = ctk.StringVar(value="Nenhuma imagem")
        ctk.CTkLabel(frame, textvariable=self.img_files_var).grid(row=1, column=1, sticky="w")

        self.img_single_pdf = ctk.CTkCheckBox(frame, text="Gerar 1 PDF único com todas as imagens", command=lambda: None)
        self.img_single_pdf.grid(row=2, column=0, padx=12, pady=6, sticky="w")
        self.img_outdir = ctk.StringVar(value=str(Path.home()))
        ctk.CTkButton(frame, text="Pasta de saída", command=self.select_img_outdir).grid(row=3, column=0, padx=12, pady=6, sticky="w")
        ctk.CTkLabel(frame, textvariable=self.img_outdir).grid(row=3, column=1, sticky="w")

        self.img_progress = ctk.CTkProgressBar(frame)
        self.img_progress.grid(row=4, column=0, columnspan=2, sticky="ew", padx=12, pady=(8,2))
        ctk.CTkButton(frame, text="Converter imagens", command=self.start_images).grid(row=5, column=0, padx=12, pady=12, sticky="w")

    def select_image_files(self):
        paths = filedialog.askopenfilenames(title="Selecionar imagens", filetypes=[("Images","*.png;*.jpg;*.jpeg;*.tiff;*.bmp")])
        if paths:
            self.img_files = list(paths)
            self.img_files_var.set(f"{len(self.img_files)} imagens selecionadas")

    def select_img_outdir(self):
        d = filedialog.askdirectory(title="Pasta de saída")
        if d:
            self.img_outdir.set(d)

    def start_images(self):
        files = getattr(self, 'img_files', [])
        if not files:
            messagebox.showwarning("Aviso", "Nenhuma imagem selecionada")
            return
        outdir = Path(self.img_outdir.get())
        outdir.mkdir(parents=True, exist_ok=True)
        single = bool(self.img_single_pdf.get())
        t = threading.Thread(target=self._do_images, args=(files, outdir, single), daemon=True)
        t.start()

    def _do_images(self, files, outdir: Path, single: bool):
        self.set_status("Convertendo imagens...")
        self.img_progress.set(0)
        try:
            outputs = images_to_pdfs(files, str(outdir), single_pdf=single)
            self.set_status(f"Imagens convertidas: {len(outputs)} arquivos gerados")
            self.img_progress.set(1.0)
        except Exception as e:
            messagebox.showerror("Erro", str(e))
            self.set_status("Erro ao converter imagens")

    # ======================= NOVA TELA: MISTURAR IMAGENS E PDFs =========================================================
    def open_tab_mix_media(self):
        self.clear_main()
        frame = ctk.CTkFrame(self.main)
        frame.pack(expand=True, pady=40)
        self.current_frame = frame

        ctk.CTkLabel(
            frame,
            text="🧩 Combinar Imagens e PDFs",
            font=ctk.CTkFont(size=22, weight="bold"),
        ).pack(pady=(10, 5))

        ctk.CTkLabel(
            frame,
            text=(
                "Selecione arquivos de imagem (PNG, JPG, JPEG, BMP, TIFF, etc.) e PDFs.\n"
                "As imagens serão convertidas em PDF automaticamente e tudo será unido em um único arquivo."
            ),
            font=ctk.CTkFont(size=13),
            justify="center",
            text_color="#aaa",
        ).pack(pady=(5, 15))

        ctk.CTkButton(
            frame,
            text="📂 Selecionar arquivos e combinar",
            command=self.mix_media_to_pdf,
            width=300,
            height=50,
            font=ctk.CTkFont(size=14, weight="bold"),
        ).pack(pady=10)


    # ======================= FUNÇÃO PRINCIPAL ================================================================================

    def mix_media_to_pdf(self):
        import tempfile
        from PyPDF2 import PdfMerger
        from PIL import Image

        filetypes = [
            ("Arquivos PDF e Imagem", "*.pdf *.png *.jpg *.jpeg *.bmp *.tiff"),
            ("Todos os arquivos", "*.*"),
        ]
        paths = filedialog.askopenfilenames(
            title="Selecione imagens e/ou PDFs", filetypes=filetypes
        )
        if not paths:
            return

        temp_pdfs = []
        try:
            with tempfile.TemporaryDirectory() as tmpdir:
                tmpdir = Path(tmpdir)
                # 1) Converter imagens em PDFs temporários
                for p in paths:
                    ext = Path(p).suffix.lower()
                    if ext in [".png", ".jpg", ".jpeg", ".bmp", ".tiff", ".tif"]:
                        try:
                            img = Image.open(p).convert("RGB")
                            pdf_temp = tmpdir / (Path(p).stem + "_conv.pdf")
                            img.save(pdf_temp)
                            temp_pdfs.append(pdf_temp)
                        except Exception as e:
                            messagebox.showwarning("Aviso", f"Erro ao converter {p}:\n{e}")
                    elif ext == ".pdf":
                        temp_pdfs.append(Path(p))

                # 2) Juntar todos os PDFs (convertidos + originais)
                if not temp_pdfs:
                    messagebox.showinfo("Aviso", "Nenhum arquivo válido foi selecionado.")
                    return

                output_path = filedialog.asksaveasfilename(
                    title="Salvar PDF combinado como...",
                    defaultextension=".pdf",
                    filetypes=[("Arquivo PDF", "*.pdf")],
                    initialfile="Resultado_Combinado.pdf",
                )
                if not output_path:
                    return

                merger = PdfMerger()
                for pdf in temp_pdfs:
                    merger.append(str(pdf))
                merger.write(output_path)
                merger.close()

                messagebox.showinfo("Sucesso", f"Arquivo criado:\n{output_path}")

        except Exception as e:
            messagebox.showerror("Erro", f"Falha ao combinar arquivos:\n{e}")
        

    # ======================= Word para PDF Tab ================================================================================================
    def open_tab_word(self):
        self.clear_main()
        frame = ctk.CTkFrame(self.main)
        frame.grid(row=0, column=0, sticky="nswe")
        self.current_frame = frame

        lbl = ctk.CTkLabel(frame, text="Converter arquivos .docx em PDF (massa)", font=ctk.CTkFont(size=16, weight="bold"))
        lbl.grid(row=0, column=0, sticky="w", padx=12, pady=8)

        btn_select = ctk.CTkButton(frame, text="Selecionar .docx", command=self.select_docx_files)
        btn_select.grid(row=1, column=0, padx=12, pady=6, sticky="w")
        self.docx_files_var = ctk.StringVar(value="Nenhum arquivo")
        ctk.CTkLabel(frame, textvariable=self.docx_files_var).grid(row=1, column=1, sticky="w")

        self.docx_outdir = ctk.StringVar(value=str(Path.home()))
        ctk.CTkButton(frame, text="Pasta de saída", command=self.select_docx_outdir).grid(row=2, column=0, padx=12, pady=6, sticky="w")
        ctk.CTkLabel(frame, textvariable=self.docx_outdir).grid(row=2, column=1, sticky="w")

        self.docx_progress = ctk.CTkProgressBar(frame)
        self.docx_progress.grid(row=3, column=0, columnspan=2, sticky="ew", padx=12, pady=(8,2))
        ctk.CTkButton(frame, text="Converter .docx → PDF", command=self.start_docx).grid(row=4, column=0, padx=12, pady=12, sticky="w")

    def select_docx_files(self):
        paths = filedialog.askopenfilenames(title="Selecionar .docx", filetypes=[("Word files","*.docx")])
        if paths:
            self.docx_files = list(paths)
            self.docx_files_var.set(f"{len(self.docx_files)} arquivos selecionados")

    def select_docx_outdir(self):
        d = filedialog.askdirectory(title="Pasta de saída")
        if d:
            self.docx_outdir.set(d)

    def start_docx(self):
        files = getattr(self, 'docx_files', [])
        if not files:
            messagebox.showwarning("Aviso", "Nenhum .docx selecionado")
            return
        outdir = Path(self.docx_outdir.get())
        outdir.mkdir(parents=True, exist_ok=True)
        t = threading.Thread(target=self._do_docx, args=(files, outdir), daemon=True)
        t.start()

    def _do_docx(self, files, outdir: Path):
        from docx import Document
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfgen import canvas

        self.set_status("Convertendo .docx em PDF...")
        self.docx_progress.set(0)

        try:
            for i, file in enumerate(files, start=1):
                try:
                    # Nome do arquivo de saída
                    out_file = outdir / (Path(file).stem + ".pdf")

                    # Lê o arquivo DOCX
                    doc = Document(file)

                    # Cria o PDF com ReportLab
                    c = canvas.Canvas(str(out_file), pagesize=A4)
                    width, height = A4
                    x, y = 50, height - 50  # Margens

                    for paragraph in doc.paragraphs:
                        text = paragraph.text.strip()
                        if not text:
                            y -= 20
                            continue
                        lines = []
                        # Quebra linhas longas automaticamente
                        while len(text) > 100:
                            split_index = text[:100].rfind(" ")
                            if split_index == -1:
                                split_index = 100
                            lines.append(text[:split_index])
                            text = text[split_index:].strip()
                        lines.append(text)
                        for line in lines:
                            c.drawString(x, y, line)
                            y -= 15
                            if y < 50:
                                c.showPage()
                                y = height - 50

                        y -= 10

                    c.save()
                except Exception as e:
                    print(f"Erro ao converter {file}: {e}")

                self.docx_progress.set(i / len(files))

            self.docx_progress.set(1.0)
            self.set_status("Conversão .docx concluída")
            messagebox.showinfo("Sucesso", f"Arquivos convertidos com sucesso em:\n{outdir}")

        except Exception as e:
            messagebox.showerror("Erro", str(e))
            self.set_status("Erro ao converter .docx")
        
    # ======================= Combinar Imagens e PDfs e Organizar ====================================================================
    def start_combine_and_organize_ui(self):
        import tempfile
        from pathlib import Path
        import tkinter as tk
        from tkinter import filedialog, messagebox
        import customtkinter as ctk
        from PIL import Image
        from PyPDF2 import PdfMerger

        # --- Janela principal ---
        win = ctk.CTkToplevel(self)
        win.title("Combinar e Organizar PDFs")
        win.geometry("700x400")

        frame = ctk.CTkFrame(win)
        frame.pack(expand=True, fill="both", padx=20, pady=20)

        title = ctk.CTkLabel(
            frame,
            text="Combinar e Organizar PDFs",
            font=ctk.CTkFont(size=22, weight="bold")
        )
        title.pack(pady=(10, 5))

        instructions = """\
    Selecione arquivos de imagem (JPG, PNG) e/ou PDFs para combinar.
    O programa irá gerar um único PDF e abrir a tela de organização
    para você rearranjar as páginas antes de salvar o resultado final."""
        ctk.CTkLabel(frame, text=instructions, justify="left", wraplength=600).pack(pady=10)

        # --- Ação principal ---
        def combine_and_open():
            try:
                # Selecionar arquivos
                files = filedialog.askopenfilenames(
                    title="Selecionar imagens e PDFs",
                    filetypes=[("Arquivos suportados", "*.pdf;*.jpg;*.jpeg;*.png")]
                )
                if not files:
                    return

                # PDF temporário
                temp_pdf_path = Path(tempfile.gettempdir()) / "combined_temp.pdf"
                merger = PdfMerger()

                for f in files:
                    ext = Path(f).suffix.lower()
                    if ext in [".jpg", ".jpeg", ".png"]:
                        img = Image.open(f).convert("RGB")
                        img_temp = Path(tempfile.gettempdir()) / f"temp_img_{Path(f).stem}.pdf"
                        img.save(img_temp)
                        merger.append(str(img_temp))
                    elif ext == ".pdf":
                        merger.append(f)

                merger.write(temp_pdf_path)
                merger.close()

                messagebox.showinfo("Sucesso", f"Arquivos combinados em:\n{temp_pdf_path}")

                # Abre automaticamente a tela de organizar com o PDF criado
                win.destroy()
                self.start_organize_ui_custom(file_path=str(temp_pdf_path))

            except Exception as e:
                messagebox.showerror("Erro", f"Ocorreu um erro: {e}")

        ctk.CTkButton(frame, text="Selecionar e Combinar", command=combine_and_open, height=40, width=200).pack(pady=30)

        ctk.CTkLabel(frame, text="Após combinar, a tela de organização abrirá automaticamente.").pack(pady=(10, 0))

    def open_tab_mix_and_organize(self):
        self.clear_main()
        frame = ctk.CTkFrame(self.main)
        frame.pack(fill="both", expand=True)
        self.current_frame = frame

        ctk.CTkLabel(
            frame,
            text="🔁 Juntar + Organizar PDFs e Imagens",
            font=ctk.CTkFont(size=22, weight="bold")
        ).pack(pady=(30, 10))

        ctk.CTkLabel(
            frame,
            text=(
                "Esta função combina automaticamente todos os arquivos PDF e imagens\n"
                "selecionados em um único arquivo e, em seguida,\n"
                "abre a tela de organização de páginas."
            ),
            justify="center",
            font=ctk.CTkFont(size=14)
        ).pack(pady=20)

        ctk.CTkButton(
            frame,
            text="Executar Juntar + Organizar",
            height=50,
            width=280,
            font=ctk.CTkFont(size=15, weight="bold"),
            command=self.run_mix_and_organize,
        ).pack(pady=30)


    def run_mix_and_organize(self):
        import fitz, os, shutil, tempfile, uuid
        from PIL import Image
        from tkinter import filedialog, messagebox

        filetypes = [
            ("PDFs e Imagens", "*.pdf *.jpg *.jpeg *.png"),
            ("Todos os arquivos", "*.*")
        ]
        files = filedialog.askopenfilenames(title="Selecione PDFs e Imagens", filetypes=filetypes)
        if not files:
            return

        try:
            # saída final (padrão)
            save_dir = os.path.expanduser("~/Pictures/Saved Pictures")
            os.makedirs(save_dir, exist_ok=True)
            output_pdf = os.path.join(save_dir, "Combinar.pdf")

            # combinador seguro
            combined = fitz.open()
            for file in files:
                ext = os.path.splitext(file)[1].lower()
                if ext == ".pdf":
                    with fitz.open(file) as sub:
                        combined.insert_pdf(sub)
                elif ext in [".jpg", ".jpeg", ".png", ".bmp", ".tif", ".tiff"]:
                    img = Image.open(file).convert("RGB")
                    tmp_img_pdf = os.path.join(tempfile.gettempdir(), f"img_{uuid.uuid4().hex}.pdf")
                    img.save(tmp_img_pdf, "PDF", resolution=100.0)
                    with fitz.open(tmp_img_pdf) as sub:
                        combined.insert_pdf(sub)
                    try:
                        os.remove(tmp_img_pdf)
                    except OSError:
                        pass
                else:
                    # ignora formatos não conhecidos
                    print("Ignorado:", file)

            combined.save(output_pdf)
            combined.close()

            # cria cópia isolada para edição — evita locks quando o arquivo original é usado por outro app
            temp_copy = os.path.join(tempfile.gettempdir(), f"organize_{uuid.uuid4().hex}.pdf")
            shutil.copy2(output_pdf, temp_copy)

            messagebox.showinfo("Sucesso", f"Arquivos combinados em:\n{output_pdf}")

            # abre a UI de organização usando a cópia isolada
            self.start_organize_ui_custom(file_path=temp_copy)

        except PermissionError:
            messagebox.showerror("Erro", "Arquivo em uso. Feche o PDF aberto em outros programas e tente novamente.")
        except Exception as e:
            messagebox.showerror("Erro", f"Erro ao combinar arquivos:\n{e}")

    def _bind_drag_events(self, frame):
        frame.bind("<Button-1>", self._on_drag_start)
        frame.bind("<B1-Motion>", self._on_drag_motion)
        frame.bind("<ButtonRelease-1>", self._on_drag_release)
        for child in frame.winfo_children():
            child.bind("<Button-1>", self._on_drag_start)
            child.bind("<B1-Motion>", self._on_drag_motion)
            child.bind("<ButtonRelease-1>", self._on_drag_release)

    def _on_drag_start(self, event):
        widget = event.widget
        while widget and not isinstance(widget, customtkinter.CTkFrame):
            widget = widget.master
        if not widget:
            return
        self.drag_data = {
            "widget": widget,
            "y": event.y_root,
            "index": self.thumbnail_frames.index(widget)
        }

    def _on_drag_motion(self, event):
        if not hasattr(self, "drag_data"):
            return
        delta_y = event.y_root - self.drag_data["y"]
        widget = self.drag_data["widget"]
        index = self.drag_data["index"]

        widget.place_configure(rely=0, y=widget.winfo_y() + delta_y)
        self.drag_data["y"] = event.y_root

        for i, frame in enumerate(self.thumbnail_frames):
            if frame == widget:
                continue
            if abs(widget.winfo_y() - frame.winfo_y()) < 30:
                self.thumbnail_frames[index], self.thumbnail_frames[i] = self.thumbnail_frames[i], self.thumbnail_frames[index]
                self._reorder_thumbnails()
                self.drag_data["index"] = i
                break

    def _on_drag_release(self, event):
        if hasattr(self, "drag_data"):
            del self.drag_data
        self._reorder_thumbnails()

    def _reorder_thumbnails(self):
        for i, frame in enumerate(self.thumbnail_frames):
            frame.place_forget()
            frame.grid(row=i // 3, column=i % 3, padx=10, pady=10)

    def _on_mousewheel(self, event):
        if hasattr(self, "canvas_frame") and hasattr(self.canvas_frame, "canvas"):
            self.canvas_frame.canvas.yview_scroll(int(-1 * (event.delta / 40)), "units")

    def delete_selected_page(self):
        from tkinter import messagebox

        if not self._selected_thumb:
            messagebox.showwarning("Aviso", "Nenhuma página selecionada.")
            return

        if not messagebox.askyesno("Confirmar", "Excluir esta página?"):
            return

        idx = self._thumb_frames.index(self._selected_thumb)

        self._selected_thumb.destroy()
        self._thumb_frames.pop(idx)
        self._orig_page_idx.pop(idx)
        self._selected_thumb = None

        for i, frame in enumerate(self._thumb_frames):
            frame.grid_forget()
            frame.grid(row=i // 3, column=i % 3, padx=10, pady=10)

        if hasattr(self, "_organize_canvas"):
            self._organize_canvas.configure(scrollregion=self._organize_canvas.bbox("all"))      

    def start_organize_ui_custom(self, file_path: str = None, input_files: list = None):
        import fitz, tempfile, os, io, threading, time
        from tkinter import filedialog, messagebox
        from PIL import Image, ImageTk
        import customtkinter as ctk
        import tkinter as tk
        
        # se não recebeu file_path, pede/combina input_files como antes
        if not file_path:
            if not input_files:
                input_files = filedialog.askopenfilenames(
                    title="Selecione imagens e/ou PDFs",
                    filetypes=[("Arquivos de imagem/PDF", "*.pdf *.png *.jpg *.jpeg *.bmp")]
                )
                if not input_files:
                    return
            # combina em arquivo temporário (uso mkstemp para evitar handle aberto)
            fd, tmpname = tempfile.mkstemp(suffix="_combined.pdf")
            os.close(fd)
            try:
                pdf_writer = fitz.open()
                for path in input_files:
                    ext = os.path.splitext(path)[1].lower()
                    if ext == ".pdf":
                        with fitz.open(path) as d:
                            pdf_writer.insert_pdf(d)
                    else:
                        pix = fitz.Pixmap(path)
                        tmp_doc = fitz.open()
                        p = tmp_doc.new_page(width=pix.width, height=pix.height)
                        p.insert_image(p.rect, pixmap=pix)
                        pdf_writer.insert_pdf(tmp_doc)
                        tmp_doc.close()
                        pix = None
                pdf_writer.save(tmpname)
                pdf_writer.close()
                file_path = tmpname
            except Exception as e:
                try:
                    os.remove(tmpname)
                except:
                    pass
                messagebox.showerror("Erro", f"Falha ao preparar PDF:\n{e}")
                return

        def save_final_pdf():
            save_path = filedialog.asksaveasfilename(
                defaultextension=".pdf",
                filetypes=[("PDF files", "*.pdf")]
            )
            if not save_path:
                return

            try:
                with fitz.open(self._organize_pdf_path) as doc:
                    new_doc = fitz.open()
                    for orig_idx in self._orig_page_idx:
                        new_doc.insert_pdf(doc, from_page=orig_idx, to_page=orig_idx)
                    new_doc.save(save_path)
                    new_doc.close()

                messagebox.showinfo("Sucesso", f"PDF salvo em:\n{save_path}")
                org_window.destroy()

            except Exception as e:
                messagebox.showerror("Erro", f"Falha ao salvar PDF:\n{e}")
        
        # --- Janela de organização ---
        org_window = ctk.CTkToplevel(self)
        org_window.title("🧩 Organizar PDF")
        org_window.geometry("1300x800")
        org_window.configure(fg_color="#1E1E1E")
        top_bar = ctk.CTkFrame(org_window, height=50)
        top_bar.pack(side="top", fill="x", padx=10, pady=(10, 5))

        ctk.CTkButton(
            top_bar,
            text="🗑 Excluir Página",
            command=self.delete_selected_page,
            fg_color="#7F1D1D",
            hover_color="#991B1B",
            width=160
        ).pack(side="left", padx=6)

        ctk.CTkButton(
            top_bar,
            text="💾 Salvar PDF Final",
            command=save_final_pdf,
            fg_color="#2563EB",
            hover_color="#1E40AF",
            width=180
        ).pack(side="right", padx=6)


        # Left: scrollável com tk.Canvas para controle fino de scroll
        # coluna esquerda FIXA (miniaturas)
        left_container = tk.Frame(org_window, bg="#1E1E1E", width=680)
        left_container.pack(side="left", fill="y", padx=(10, 4), pady=10)
        left_container.pack_propagate(False)

        canvas = tk.Canvas(left_container, bg="#1E1E1E", highlightthickness=0)
        canvas.pack(side="left", fill="both", expand=True)

        self._organize_canvas = canvas

        vscroll = tk.Scrollbar(left_container, orient="vertical", command=canvas.yview)
        vscroll.pack(side="right", fill="y")

        inner_frame = tk.Frame(canvas, bg="#1E1E1E")
        canvas.create_window((0, 0), window=inner_frame, anchor="nw")

        def _on_configure(event):
            canvas.configure(scrollregion=canvas.bbox("all"))
        inner_frame.bind("<Configure>", _on_configure)

        # right preview
        right_frame = ctk.CTkFrame(org_window, fg_color="#111111")
        right_frame.pack(side="left", fill="both", expand=True, padx=(4, 10), pady=10)
        preview_canvas = tk.Canvas(right_frame, bg="#111111", highlightthickness=0)
        preview_canvas.pack(fill="both", expand=True, padx=6, pady=6)

        # estado local
        self._organize_pdf_path = file_path
        self._thumb_frames = []
        self._thumb_images = []
        self._orig_page_idx = []
        self._selected_thumb = None

        # scroll sensível
        def on_mousewheel(event):
            delta = -int(event.delta / 40)
            canvas.yview_scroll(delta, "units")
        canvas.bind_all("<MouseWheel>", on_mousewheel)

        # helpers para imagens
        def pil_from_page(pdfpath, page_num, dpi=90, target_width=None):
            try:
                with fitz.open(pdfpath) as d:
                    page = d.load_page(page_num - 1)
                    pix = page.get_pixmap(dpi=dpi)
                    data = pix.tobytes("png")
                img = Image.open(io.BytesIO(data)).convert("RGB")
                if target_width:
                    ratio = target_width / img.width
                    img = img.resize((int(target_width), int(img.height * ratio)), Image.LANCZOS)
                return img
            except Exception as e:
                print("[pil_from_page err]", e)
                return None

        def show_preview(page_num):
            try:
                img = pil_from_page(self._organize_pdf_path, page_num, dpi=150)
                if img is None:
                    return
                cw = preview_canvas.winfo_width() or 600
                ch = preview_canvas.winfo_height() or 800
                ratio = min(cw / img.width, ch / img.height)
                new_size = (int(img.width * ratio), int(img.height * ratio))
                img = img.resize(new_size, Image.LANCZOS)
                preview_canvas.delete("all")
                tkimg = ImageTk.PhotoImage(img)
                preview_canvas.create_image((cw - new_size[0]) // 2, (ch - new_size[1]) // 2, anchor="nw", image=tkimg)
                preview_canvas.image = tkimg
            except Exception as e:
                print("[preview err]", e)

        # clique vs arrastar
        CLICK_DRAG_THRESHOLD_SECS = 0.15
        MOVE_THRESHOLD_PX = 6
        drag = {"widget": None, "start_x": 0, "start_y": 0, "start_index": None, "press_time": 0}

        def on_thumb_press(ev, frame, orig_index):
            drag["widget"] = frame
            drag["start_x"] = ev.x_root
            drag["start_y"] = ev.y_root
            drag["start_index"] = self._thumb_frames.index(frame)
            drag["press_time"] = time.time()
            frame.lift()
            frame.configure(fg_color="#2E2E2E")
            # marca seleção visual
            if self._selected_thumb and self._selected_thumb != frame:
                self._selected_thumb.configure(fg_color="#2E2E2E")

            self._selected_thumb = frame
            frame.configure(fg_color="#555555")


        def on_thumb_motion(ev):
            w = drag.get("widget")
            if not w:
                return
            dx = ev.x_root - drag["start_x"]
            dy = ev.y_root - drag["start_y"]
            moved_enough = abs(dx) > MOVE_THRESHOLD_PX or abs(dy) > MOVE_THRESHOLD_PX
            if time.time() - drag["press_time"] < CLICK_DRAG_THRESHOLD_SECS and not moved_enough:
                return
            # move
            w.place(in_=inner_frame, x=w.winfo_x() + dx, y=w.winfo_y() + dy)
            drag["start_x"] = ev.x_root
            drag["start_y"] = ev.y_root
            # highlight
            for f in self._thumb_frames:
                if f is w: continue
                x1, y1 = f.winfo_x(), f.winfo_y()
                x2, y2 = x1 + f.winfo_width(), y1 + f.winfo_height()
                rel_x = ev.x_root - inner_frame.winfo_rootx()
                rel_y = ev.y_root - inner_frame.winfo_rooty()
                if x1 < rel_x < x2 and y1 < rel_y < y2:
                    f.configure(fg_color="#444444")
                else:
                    f.configure(fg_color="#2E2E2E")

        def on_thumb_release(ev):
            w = drag.get("widget")
            if not w:
                return
            elapsed = time.time() - drag["press_time"]
            for f in self._thumb_frames:
                f.configure(fg_color="#2E2E2E")
            moved = abs(ev.x_root - drag["start_x"]) > MOVE_THRESHOLD_PX or abs(ev.y_root - drag["start_y"]) > MOVE_THRESHOLD_PX
            if elapsed < CLICK_DRAG_THRESHOLD_SECS and not moved:
                orig_index = getattr(w, "orig_index", None)
                if orig_index is not None:
                    show_preview(orig_index + 1)
                drag.update({"widget": None})
                return
            # compute new index
            rel_x = ev.x_root - inner_frame.winfo_rootx()
            rel_y = ev.y_root - inner_frame.winfo_rooty()
            col = max(0, int(rel_x // (240 + 20)))
            row = max(0, int(rel_y // (320 + 20)))
            new_index = min(row * 3 + col, len(self._thumb_frames) - 1)
            try:
                start_idx = self._thumb_frames.index(w)
            except ValueError:
                start_idx = drag.get("start_index", 0)
            item = self._thumb_frames.pop(start_idx)
            orig = self._orig_page_idx.pop(start_idx)
            self._thumb_frames.insert(new_index, item)
            self._orig_page_idx.insert(new_index, orig)
            # re-grid
            for i, frame in enumerate(self._thumb_frames):
                frame.place_forget()
                frame.grid(row=i // 3, column=i % 3, padx=10, pady=10)

            elf._organize_canvas.configure(
                scrollregion=self._organize_canvas.bbox("all")
            )
                
            drag.update({"widget": None})
        

        # geração de thumbnails em thread
        def generate_thumbnails():
            try:
                with fitz.open(self._organize_pdf_path) as doc:
                    n = len(doc)
                    for i in range(n):
                        pil = pil_from_page(self._organize_pdf_path, i + 1, dpi=90, target_width=190)
                        if pil is None:
                            continue
                        tkthumb = ImageTk.PhotoImage(pil)
                        container = ctk.CTkFrame(inner_frame, fg_color="#2E2E2E", corner_radius=8)
                        lbl_img = tk.Label(container, image=tkthumb, bd=0, bg="#2E2E2E")
                        lbl_img.pack(padx=6, pady=6)
                        lbl_text = ctk.CTkLabel(container, text=f"P. {i+1}", fg_color="#2E2E2E", text_color="#EEE")
                        lbl_text.pack(pady=(0,6))
                        container.orig_index = i
                        container.page_num = i + 1
                        # binds
                        container.bind("<ButtonPress-1>", lambda e, fr=container, idx=i: on_thumb_press(e, fr, idx))
                        container.bind("<B1-Motion>", lambda e: on_thumb_motion(e))
                        container.bind("<ButtonRelease-1>", lambda e: on_thumb_release(e))
                        lbl_img.bind("<ButtonPress-1>", lambda e, fr=container, idx=i: on_thumb_press(e, fr, idx))
                        lbl_img.bind("<B1-Motion>", lambda e: on_thumb_motion(e))
                        lbl_img.bind("<ButtonRelease-1>", lambda e: on_thumb_release(e))
                        # store
                        self._thumb_images.append(tkthumb)
                        self._thumb_frames.append(container)
                        self._orig_page_idx.append(i)
                        idx = len(self._thumb_frames) - 1
                        container.grid(row=idx // 3, column=idx % 3, padx=10, pady=10)
                inner_frame.update_idletasks()
                self._organize_canvas.configure(
                scrollregion=self._organize_canvas.bbox("all"))
            except Exception as e:
                print("[generate_thumbnails err]", e)

        threading.Thread(target=generate_thumbnails, daemon=True).start()

        # salvar conforme ordem atual
        def save_final_pdf():
            save_path = filedialog.asksaveasfilename(defaultextension=".pdf", filetypes=[("PDF files", "*.pdf")])
            if not save_path:
                return
            try:
                with fitz.open(self._organize_pdf_path) as doc:
                    new_doc = fitz.open()
                    for orig_idx in self._orig_page_idx:
                        new_doc.insert_pdf(doc, from_page=orig_idx, to_page=orig_idx)
                    new_doc.save(save_path)
                    new_doc.close()
                messagebox.showinfo("Sucesso", f"PDF salvo em:\n{save_path}")
                org_window.destroy()
            except Exception as e:
                messagebox.showerror("Erro", f"Falha ao salvar PDF:\n{e}")

#================================ Renomeador de Arquivos ============================================
    def open_tab_renamer(self):
        self.clear_main()
        frame = ctk.CTkFrame(self.main)
        frame.grid(row=0, column=0, sticky="nsew", padx=20, pady=20)
        self.current_frame = frame

        # ===== Título =====
        ctk.CTkLabel(
            frame, text="Renomeador de PDFs", 
            font=ctk.CTkFont(size=18, weight="bold")
        ).grid(row=0, column=0, columnspan=4, pady=(5, 15))

        # ===== Listas =====
        ctk.CTkLabel(frame, text="Lista Principal:").grid(row=1, column=0, sticky="w")
        self.main_list_text = ctk.CTkTextbox(frame, height=80, width=250)
        self.main_list_text.grid(row=2, column=0, padx=(0, 10), pady=(0, 10))

        ctk.CTkLabel(frame, text="Lista Auxiliar (Opcional):").grid(row=1, column=1, sticky="w")
        self.aux_list_text = ctk.CTkTextbox(frame, height=80, width=250)
        self.aux_list_text.grid(row=2, column=1, padx=(0, 10), pady=(0, 10))

        # ===== Prefixo, sufixo, recortes =====
        options_frame = ctk.CTkFrame(frame)
        options_frame.grid(row=3, column=0, columnspan=2, pady=(5, 10), sticky="ew")

        ctk.CTkLabel(options_frame, text="Prefixo:").grid(row=0, column=0, padx=5, pady=5)
        self.prefix_entry = ctk.CTkEntry(options_frame, width=120)
        self.prefix_entry.grid(row=0, column=1, padx=5, pady=5)

        ctk.CTkLabel(options_frame, text="Sufixo:").grid(row=0, column=2, padx=5, pady=5)
        self.suffix_entry = ctk.CTkEntry(options_frame, width=120)
        self.suffix_entry.grid(row=0, column=3, padx=5, pady=5)

        ctk.CTkLabel(options_frame, text="Manter primeiros X:").grid(row=1, column=0, padx=5, pady=5)
        self.keep_first_entry = ctk.CTkEntry(options_frame, width=80)
        self.keep_first_entry.insert(0, "0")
        self.keep_first_entry.grid(row=1, column=1, padx=5, pady=5)

        ctk.CTkLabel(options_frame, text="Manter últimos X:").grid(row=1, column=2, padx=5, pady=5)
        self.keep_last_entry = ctk.CTkEntry(options_frame, width=80)
        self.keep_last_entry.insert(0, "0")
        self.keep_last_entry.grid(row=1, column=3, padx=5, pady=5)

        # ===== Botões =====
        btn_frame = ctk.CTkFrame(frame)
        btn_frame.grid(row=4, column=0, columnspan=2, pady=(10, 10))

        ctk.CTkButton(
            btn_frame, text="Selecionar PDFs", 
            command=self.select_pdfs_for_rename, width=160
        ).grid(row=0, column=0, padx=10, pady=5)

        ctk.CTkButton(
            btn_frame, text="Renomear Arquivos", 
            command=self.start_rename_process, width=160
        ).grid(row=0, column=1, padx=10, pady=5)

        # ===== Pré-visualização e progresso =====
        ctk.CTkLabel(frame, text="Pré-visualização dos novos nomes:").grid(
            row=5, column=0, columnspan=2, sticky="w", pady=(5, 0)
        )
        self.preview_box = ctk.CTkTextbox(frame, height=160, width=550)
        self.preview_box.grid(row=6, column=0, columnspan=2, padx=5, pady=5)

        self.rename_progress = ctk.CTkProgressBar(frame)
        self.rename_progress.grid(row=7, column=0, columnspan=2, sticky="ew", pady=(5, 0))
        self.rename_progress.set(0)

        # ===== Variáveis de controle =====
        self.pdf_files_to_rename = []
        self.preview_box.insert("1.0", "Nenhum arquivo selecionado ainda.")

    # ===== Selecionar arquivos =====
    def select_pdfs_for_rename(self):
        files = filedialog.askopenfilenames(title="Selecionar PDFs", filetypes=[("PDF files","*.pdf")])
        if not files:
            return
        self.pdf_files_to_rename = list(files)
        self.preview_box.delete("1.0", "end")
        self.preview_box.insert("1.0", "Gerando pré-visualização...\n")
        threading.Thread(target=self.gen_preview, daemon=True).start()

    def gen_preview(self):
        """Gera a pré-visualização dos nomes finais (executa em thread)."""
        import pytesseract, fitz
        from pathlib import Path
        from rapidfuzz import fuzz

        main_list = [x.strip() for x in self.main_list_text.get("1.0", "end").splitlines() if x.strip()]
        aux_list = [x.strip() for x in self.aux_list_text.get("1.0", "end").splitlines() if x.strip()]
        prefix = self.prefix_entry.get().strip()
        suffix = self.suffix_entry.get().strip()

        try:
            keep_first = int(self.keep_first_entry.get() or 0)
        except Exception:
            keep_first = 0
        try:
            keep_last = int(self.keep_last_entry.get() or 0)
        except Exception:
            keep_last = 0

        files = list(self.pdf_files_to_rename)
        total = len(files)
        if total == 0:
            return

        # mostrar progresso de leitura
        self.rename_progress.set(0)
        self.preview_box.delete("1.0", "end")
        self.preview_box.insert("1.0", "Gerando pré-visualização (lendo arquivos)...\n")

        preview_lines = []
        extracted_text_cache = []

        for i, file in enumerate(files, start=1):
            try:
                txt = ""
                with fitz.open(file) as doc:
                    for page in doc:
                        text = page.get_text("text")
                        txt += text or ""
                        # se pouco texto, tenta OCR na página
                        if len(txt.strip()) < 30:
                            pix = page.get_pixmap(dpi=200)
                            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                            try:
                                txt += pytesseract.image_to_string(img, lang="por", config="--psm 6")
                            except Exception:
                                # fallback: sem OCR se pytesseract faltar/config falhar
                                pass

                extracted_text_cache.append(txt)
                # Encontrar matches
                match_main = self._find_best_match(txt, main_list)
                match_aux = self._find_best_match(txt, aux_list)

                orig_stem = Path(file).stem

                # Constrói o novo nome DO ZERO (não manter o nome antigo, salvo os recortes)
                parts = []

                # prefixo primeiro (se houver)
                if prefix:
                    parts.append(prefix)

                # manter primeiros X caracteres do nome original (opcional)
                if keep_first > 0:
                    parts.append(orig_stem[:keep_first])

                # adicionar match principal / auxiliar
                if match_main:
                    parts.append(match_main)
                if match_aux:
                    parts.append(match_aux)

                # manter últimos X caracteres (após matches ou onde fizer sentido)
                if keep_last > 0:
                    parts.append(orig_stem[-keep_last:])

                # sufixo por último
                if suffix:
                    parts.append(suffix)

                # remove partes vazias e normaliza underscores
                parts = [p for p in parts if p and str(p).strip()]
                if parts:
                    new_name = "_".join(parts)
                else:
                    # se nada escolhido, usa somente prefix/sufixo se existirem, senão mantém nome original
                    if prefix or suffix:
                        middle = orig_stem if (keep_first > 0 or keep_last > 0) else ""
                        new_name = "_".join([p for p in [prefix, middle, suffix] if p])
                    else:
                        new_name = orig_stem  # fallback: mantém o nome original

                # normalização final: remover espaços duplos, proibidos, e uppercase opcional
                new_name = re.sub(r'\s+', '_', new_name).strip("_")
                new_name = f"{new_name}.pdf"

                preview_lines.append((file, new_name))

                # atualiza preview box incrementalmente
                self.preview_box.insert("end", f"{Path(file).name}  →  {new_name}\n")

            except Exception as e:
                self.preview_box.insert("end", f"Erro ao ler {Path(file).name}: {e}\n")
            # atualiza barra de leitura
            self.rename_progress.set(i / total * 0.5)  # leitura = 50% do processo
        # fim loop

        # guardamos preview_lines para renomeação real
        self._preview_lines = preview_lines
        # marca progresso leitura concluída (50%)
        self.rename_progress.set(0.5)
        self.preview_box.insert("end", "\nPré-visualização gerada. Clique em 'Renomear Arquivos' para aplicar.\n")

    def _find_best_match(self, text, candidates):
        """Busca o melhor candidato usando Partial Ratio (rapidfuzz) com threshold."""
        try:
            from rapidfuzz import fuzz
        except Exception:
            import difflib
            # fallback simples
            best = None
            best_score = 0.0
            for c in candidates:
                score = difflib.SequenceMatcher(None, c.lower(), text.lower()).ratio() * 100
                if score > best_score and score >= 70:
                    best = c
                    best_score = score
            return best

        best = None
        best_score = 0.0
        for c in candidates:
            if not c:
                continue
            score = fuzz.partial_ratio(c.lower(), text.lower())
            if score > best_score and score >= 70:
                best = c
                best_score = score
        return best

    def start_rename_process(self):
        """Inicia thread que efetivamente copia/renomeia os arquivos usando preview salvo."""
        if not getattr(self, "_preview_lines", None):
            messagebox.showwarning("Aviso", "Gere a pré-visualização antes de renomear (Selecione PDFs).")
            return
        # pede diretório de saída
        outdir = filedialog.askdirectory(title="Escolher pasta de saída")
        if not outdir:
            return
        # start thread
        threading.Thread(target=self._rename_files, args=(outdir,), daemon=True).start()

    def _rename_files(self, outdir):
        """Copia arquivos para outdir com os nomes mostrados na pré-visualização.
           Lida com conflitos (adiciona contador)."""
        import shutil
        from pathlib import Path

        lines = getattr(self, "_preview_lines", [])
        total = len(lines)
        if total == 0:
            return

        outdir = Path(outdir)
        outdir.mkdir(parents=True, exist_ok=True)

        # progress: leitura já corresponde a primeira metade; faremos a renomeação na segunda metade
        self.rename_progress.set(0.5)
        renamed = 0

        for i, (src, new_name) in enumerate(lines, start=1):
            try:
                base = Path(new_name).stem
                dest = outdir / f"{base}.pdf"
                # evitar sobrescrever: acrescenta (1),(2)...
                counter = 1
                while dest.exists():
                    dest = outdir / f"{base}({counter}).pdf"
                    counter += 1
                shutil.copy2(src, dest)
                renamed += 1
            except Exception as e:
                print("Erro ao copiar/renomear:", src, e)
            # atualiza barra (50%..100%)
            self.rename_progress.set(0.5 + (i / total) * 0.5)

        messagebox.showinfo("Concluído", f"Operação finalizada: {renamed} arquivo(s) processado(s).\nDestino: {outdir}")
        self.rename_progress.set(1.0)

#======================= Dividir PDF =========================================================================================
    def open_tab_split_pdf(self):
        self.clear_main()

        frame = ctk.CTkFrame(self.main)
        frame.pack(fill="both", expand=True)
        self.current_frame = frame

        ctk.CTkLabel(
            frame,
            text="✂️ Dividir PDF",
            font=ctk.CTkFont(size=22, weight="bold")
        ).pack(pady=(30, 10))

        ctk.CTkLabel(
            frame,
            text=(
                "Divida arquivos PDF em vários arquivos menores\n"
            ),
            justify="center",
            font=ctk.CTkFont(size=14)
        ).pack(pady=20)

        ctk.CTkButton(
            frame,
            text="Selecionar PDF",
            height=50,
            width=280,
            font=ctk.CTkFont(size=15, weight="bold"),
            command=self.run_split_pdf,
        ).pack(pady=30)


    def run_split_pdf(self):
        from tkinter import filedialog

        file_path = filedialog.askopenfilename(
            title="Selecione um PDF",
            filetypes=[("PDF", "*.pdf")]
        )
        if not file_path:
            return

        self.start_visual_pdf_ui(file_path=file_path, mode="split")   

    def start_visual_pdf_ui(self, file_path: str, mode: str = "organize"):
        import fitz, os, io, threading, time
        from tkinter import filedialog, messagebox
        from PIL import Image, ImageTk
        import customtkinter as ctk
        import tkinter as tk

        # ================== Janela ==================
        win = ctk.CTkToplevel(self)
        win.geometry("1300x800")
        win.title("🧩 Organizar PDF" if mode == "organize" else "✂️ Dividir PDF")
        win.configure(fg_color="#1E1E1E")

        def on_close():
            try:
                self._pdf_doc.close()
            except:
                pass
            win.destroy()

        win.protocol("WM_DELETE_WINDOW", on_close)


        # ================== Top bar ==================
        top_bar = ctk.CTkFrame(win, height=60)
        top_bar.pack(side="top", fill="x", padx=10, pady=(10, 5))

        if mode == "split":
            split_mode = tk.StringVar(value="interval")

            ctk.CTkRadioButton(top_bar, text="Intervalos",
                variable=split_mode, value="interval").pack(side="left", padx=6)
            ctk.CTkRadioButton(top_bar, text="Páginas individuais",
                variable=split_mode, value="single").pack(side="left", padx=6)
            ctk.CTkRadioButton(top_bar, text="A cada X páginas",
                variable=split_mode, value="every").pack(side="left", padx=6)

            interval_entry = ctk.CTkEntry(top_bar, placeholder_text="Ex: 1-3,5,7-9", width=200)
            interval_entry.pack(side="left", padx=10)

            every_entry = ctk.CTkEntry(top_bar, placeholder_text="X páginas", width=120)
            every_entry.pack(side="left", padx=6)

            ctk.CTkButton(
                top_bar,
                text="🗑 Excluir Página",
                width=160,
                fg_color="#7F1D1D",
                hover_color="#991B1B",
                command=lambda: delete_selected_page()
            ).pack(side="right", padx=6)

            ctk.CTkButton(
                top_bar,
                text="✂️ Dividir PDF",
                width=160,
                fg_color="#2563EB",
                hover_color="#1E40AF",
                command=lambda: self.execute_split_visual(
                    split_mode.get(),
                    interval_entry.get(),
                    every_entry.get()
                )
            ).pack(side="right", padx=6)

        else:
            ctk.CTkButton(
                top_bar,
                text="💾 Salvar PDF",
                width=160,
                fg_color="#16A34A",
                hover_color="#166534",
                command=lambda: self.save_organized_pdf(win)
            ).pack(side="right", padx=6)

        # ================== Layout ==================
        left = tk.Frame(win, bg="#1E1E1E", width=680)
        left.pack(side="left", fill="y", padx=(10, 4), pady=10)
        left.pack_propagate(False)

        canvas = tk.Canvas(left, bg="#1E1E1E", highlightthickness=0)
        canvas.pack(side="left", fill="both", expand=True)

        scrollbar = tk.Scrollbar(left, orient="vertical", command=canvas.yview)
        scrollbar.pack(side="right", fill="y", padx=(6, 0))
        
        canvas.configure(yscrollcommand=scrollbar.set)

        inner = tk.Frame(canvas, bg="#1E1E1E")
        drag_indicator = tk.Frame(inner, height=4, bg="#3B82F6")
        drag_indicator.place_forget()
        canvas.create_window((0, 0), window=inner, anchor="nw")

        def _on_cfg(_):
            canvas.configure(scrollregion=canvas.bbox("all"))
        inner.bind("<Configure>", _on_cfg)

        right = ctk.CTkFrame(win, fg_color="#111111")
        right.pack(side="left", fill="both", expand=True, padx=(4, 10), pady=10)

        preview = tk.Canvas(right, bg="#111111", highlightthickness=0)
        preview.pack(fill="both", expand=True, padx=6, pady=6)

        def on_mousewheel(event):
            canvas.yview_scroll(-1 * int(event.delta / 120), "units")

        canvas.bind_all("<MouseWheel>", on_mousewheel)

        # ================== Estado ==================
        self._organize_pdf_path = file_path
        
        # inicializa se não existir
        if not hasattr(self, "_thumb_frames"):
            self._thumb_frames = []
        if not hasattr(self, "_thumb_images"):
            self._thumb_images = []
        if not hasattr(self, "_orig_page_idx"):
            self._orig_page_idx = []

        # limpa estado
        self._thumb_frames.clear()
        self._thumb_images.clear()
        self._orig_page_idx.clear()
        self._selected_thumb = None

        self._pdf_doc = fitz.open(self._organize_pdf_path)

        # ================== Helpers ==================
        def pil_from_page(page):
            with fitz.open(self._organize_pdf_path) as d:
                p = d.load_page(page)
                pix = p.get_pixmap(dpi=90)
            return Image.open(io.BytesIO(pix.tobytes("png"))).convert("RGB")

        def show_preview(idx):
            img = pil_from_page(idx)
            w = preview.winfo_width() or 600
            h = preview.winfo_height() or 800
            r = min(w / img.width, h / img.height)
            img = img.resize((int(img.width * r), int(img.height * r)), Image.LANCZOS)
            preview.delete("all")
            tkimg = ImageTk.PhotoImage(img)
            preview.create_image((w - img.width) // 2, (h - img.height) // 2,
                                 anchor="nw", image=tkimg)
            preview.image = tkimg

        def delete_selected_page():
            if not self._selected_thumb:
                messagebox.showwarning("Aviso", "Nenhuma página selecionada.")
                return

            idx = self._thumb_frames.index(self._selected_thumb)

            self._selected_thumb.destroy()
            self._thumb_frames.pop(idx)
            self._orig_page_idx.pop(idx)
            self._selected_thumb = None

            for i, frame in enumerate(self._thumb_frames):
                frame.grid_forget()
                frame.grid(row=i // 3, column=i % 3, padx=10, pady=10)

            canvas.configure(scrollregion=canvas.bbox("all"))    

        # ================== Thumbnails ==================
        def load_thumbs():
            with fitz.open(self._organize_pdf_path) as doc:
                for orig_idx in range(len(doc)):
                    img = pil_from_page(orig_idx)
                    img.thumbnail((190, 270))
                    tkimg = ImageTk.PhotoImage(img)

                    box = ctk.CTkFrame(inner, fg_color="#2E2E2E", corner_radius=8)
                    box.orig_index = orig_idx

                    img_label = tk.Label(box, image=tkimg, bg="#2E2E2E")
                    img_label.pack(padx=6, pady=6)

                    text_label = ctk.CTkLabel(box, text=f"P. {orig_idx + 1}")
                    text_label.pack(pady=(0, 6))

                    def on_thumb_click(event, frame=box):
                        if self._selected_thumb:
                            self._selected_thumb.configure(fg_color="#2E2E2E")
                        self._selected_thumb = frame
                        frame.configure(fg_color="#3B82F6")
                        show_preview(frame.orig_index)

                    def on_thumb_press(event, frame=box):
                        self._dragged_frame = frame
                        self._drag_start_index = self._thumb_frames.index(frame)
                        frame.configure(fg_color="#2563EB")

                    def on_thumb_drag(event, frame=box):
                        if self._dragged_frame != frame:
                            return

                        y = canvas.winfo_pointery() - inner.winfo_rooty()

                        # altura média do item
                        item_h = frame.winfo_height() + 20

                        target_idx = int(y // item_h) * 3
                        x = canvas.winfo_pointerx() - inner.winfo_rootx()
                        y = canvas.winfo_pointery() - inner.winfo_rooty()

                        item_w = frame.winfo_width() + 20
                        item_h = frame.winfo_height() + 20

                        col = int(x // item_w)
                        row = int(y // item_h)

                        col = max(0, min(col, 2))  # 3 colunas: 0,1,2
                        target_idx = row * 3 + col
                        target_idx = max(0, min(target_idx, len(self._thumb_frames)))

                        # posição visual do indicador
                        if target_idx < len(self._thumb_frames):
                            target = self._thumb_frames[target_idx]
                            drag_indicator.place(
                                x=target.winfo_x(),
                                y=target.winfo_y() - 6,
                                width=target.winfo_width()
                            )
                        else:
                            last = self._thumb_frames[-1]
                            drag_indicator.place(
                                x=last.winfo_x(),
                                y=last.winfo_y() + last.winfo_height() + 6,
                                width=last.winfo_width()
                            )
                            
                        self._drop_target_index = target_idx

                    def on_thumb_release(event, frame=box):
                        drag_indicator.place_forget()

                        if not hasattr(self, "_drop_target_index"):
                            return

                        src = self._drag_start_index
                        tgt = self._drop_target_index

                        if src != tgt and tgt < len(self._thumb_frames):
                            self._thumb_frames.insert(tgt, self._thumb_frames.pop(src))
                            self._orig_page_idx.insert(tgt, self._orig_page_idx.pop(src))

                        # regrid completo
                        for i, f in enumerate(self._thumb_frames):
                            f.grid_forget()
                            f.grid(row=i // 3, column=i % 3, padx=10, pady=10)
                            f.configure(fg_color="#2E2E2E")

                        frame.configure(fg_color="#3B82F6")
                        self._selected_thumb = frame

                        del self._drop_target_index

                    for w in (box, img_label, text_label):
                        w.bind("<Button-1>", on_thumb_click)
                        w.bind("<ButtonPress-1>", on_thumb_press)
                        w.bind("<B1-Motion>", on_thumb_drag)
                        w.bind("<ButtonRelease-1>", on_thumb_release)

                    # mantém referência da imagem (evita GC)
                    self._thumb_images.append(tkimg)

                    # registra estado
                    self._thumb_frames.append(box)
                    self._orig_page_idx.append(orig_idx)

                    # posiciona no grid
                    idx = len(self._thumb_frames) - 1
                    box.grid(
                        row=idx // 3,
                        column=idx % 3,
                        padx=10,
                        pady=10
                    )    

                canvas.configure(scrollregion=canvas.bbox("all"))

        threading.Thread(target=load_thumbs, daemon=True).start()
      

    def execute_split_visual(self, mode, interval_text, every_text):
        import fitz
        from tkinter import filedialog, messagebox
        import os

        if not self._orig_page_idx:
            messagebox.showerror("Erro", "Nenhuma página carregada.")
            return

        save_dir = filedialog.askdirectory(title="Escolha a pasta de destino")
        if not save_dir:
            return

        try:
            with fitz.open(self._organize_pdf_path) as doc:
                ordered_pages = self._orig_page_idx[:]  # ORDEM VISUAL REAL

                def save_part(page_indexes, part_idx):
                    new_doc = fitz.open()
                    for p in page_indexes:
                        new_doc.insert_pdf(doc, from_page=p, to_page=p)
                    out = os.path.join(save_dir, f"parte_{part_idx}.pdf")
                    new_doc.save(out)
                    new_doc.close()

                # =========================
                # PÁGINAS INDIVIDUAIS
                # =========================
                if mode == "single":
                    for i, p in enumerate(ordered_pages, start=1):
                        save_part([p], i)

                # =========================
                # A CADA X PÁGINAS
                # =========================
                elif mode == "every":
                    try:
                        step = int(every_text)
                        if step <= 0:
                            raise ValueError
                    except:
                        messagebox.showerror("Erro", "Informe um número válido.")
                        return

                    idx = 1
                    for i in range(0, len(ordered_pages), step):
                        save_part(ordered_pages[i:i + step], idx)
                        idx += 1

                # =========================
                # INTERVALOS (VISUAIS)
                # =========================
                elif mode == "interval":
                    try:
                        parts = interval_text.split(",")
                        idx = 1
                        for part in parts:
                            part = part.strip()
                            if "-" in part:
                                a, b = map(int, part.split("-"))
                                pages = ordered_pages[a - 1:b]
                            else:
                                pages = [ordered_pages[int(part) - 1]]
                            save_part(pages, idx)
                            idx += 1
                    except Exception:
                        messagebox.showerror("Erro", "Intervalo inválido.")
                        return

            messagebox.showinfo("Concluído", "PDF dividido com sucesso.")

        except Exception as e:
            messagebox.showerror("Erro", f"Falha ao dividir PDF:\n{e}")
            

#======================= PDF para PNG e JPEG =========================================================================================
    def convert_pdfs_to_images(self, pdf_files, output_dir, img_format):
        def worker():
            for pdf_path in pdf_files:
                pdf_name = os.path.splitext(os.path.basename(pdf_path))[0]

                with fitz.open(pdf_path) as doc:
                    for i, page in enumerate(doc):
                        pix = page.get_pixmap(dpi=200)

                        ext = "png" if img_format == "png" else "jpg"
                        output_path = os.path.join(
                            output_dir,
                            f"{pdf_name}_P{i+1}.{ext}"
                        )

                        if img_format == "jpg":
                            pix.save(output_path, output="jpeg")
                        else:
                            pix.save(output_path)

            self.after(0, lambda: messagebox.showinfo(
                "Concluído",
                "Conversão finalizada com sucesso."
            ))

        threading.Thread(target=worker, daemon=True).start()
            
    def open_pdf_to_images(self):
        self.clear_main()  # mesma função que você já usa para trocar telas

        container = ctk.CTkFrame(self.main, fg_color="#1E1E1E")
        container.pack(fill="both", expand=True, padx=20, pady=20)

        self.current_frame = container

        title = ctk.CTkLabel(
            container,
            text="Converter PDFs em Imagens",
            font=("Segoe UI", 18, "bold")
        )
        title.pack(pady=(10, 30))

        format_var = ctk.StringVar(value="png")

        radio_frame = ctk.CTkFrame(container, fg_color="transparent")
        radio_frame.pack(pady=(0, 20))

        ctk.CTkRadioButton(
            radio_frame, text="PNG", variable=format_var, value="png"
        ).pack(side="left", padx=20)

        ctk.CTkRadioButton(
            radio_frame, text="JPG", variable=format_var, value="jpg"
        ).pack(side="left", padx=20)

        def select_files():
            pdf_files = filedialog.askopenfilenames(
                title="Selecionar PDFs",
                filetypes=[("PDF", "*.pdf")]
            )
            if not pdf_files:
                return

            output_dir = filedialog.askdirectory(
                title="Selecionar pasta de destino"
            )
            if not output_dir:
                return

            self.convert_pdfs_to_images(pdf_files, output_dir, format_var.get())

        btn_select = ctk.CTkButton(
            container,
            text="Selecionar PDFs",
            height=40,
            command=select_files
        )
        btn_select.pack(pady=10)
            
#======================== PDF para Word =========================================================================================================
    def run_pdf_to_word_ocr(self):
        import fitz
        import pytesseract
        from PIL import Image
        from tkinter import filedialog, messagebox
        from docx import Document
        from docx.shared import Inches
        import os
        import io

        pdfs = filedialog.askopenfilenames(
            title="Selecione PDFs",
            filetypes=[("PDF", "*.pdf")]
        )
        if not pdfs:
            return

        out_dir = filedialog.askdirectory(title="Selecione a pasta de destino")
        if not out_dir:
            return

        for pdf_path in pdfs:
            try:
                doc = fitz.open(pdf_path)
                word = Document()

                for page in doc:
                    pix = page.get_pixmap(dpi=300)
                    img = Image.open(io.BytesIO(pix.tobytes("png")))

                    # OCR com layout
                    ocr_data = pytesseract.image_to_data(
                        img,
                        lang="por",
                        output_type=pytesseract.Output.DICT,
                        config="--psm 6"
                    )

                    current_line = ""
                    last_y = None

                    for i in range(len(ocr_data["text"])):
                        text = ocr_data["text"][i].strip()
                        if not text:
                            continue

                        y = ocr_data["top"][i]

                        if last_y is not None and abs(y - last_y) > 15:
                            word.add_paragraph(current_line)
                            current_line = text
                        else:
                            current_line += " " + text

                        last_y = y

                    if current_line.strip():
                        word.add_paragraph(current_line)

                    # adiciona imagem da página como referência visual (opcional)
                    img_path = os.path.join(out_dir, "_page_tmp.png")
                    img.save(img_path)
                    word.add_picture(img_path, width=Inches(6))

                name = os.path.splitext(os.path.basename(pdf_path))[0]
                out_path = os.path.join(out_dir, f"{name}.docx")
                word.save(out_path)
                doc.close()

            except Exception as e:
                messagebox.showerror("Erro", f"Falha ao converter:\n{e}")
                return

        messagebox.showinfo("Concluído", "Conversão finalizada.")

        
    def open_tab_pdf_to_word(self):
        # ========= UI =========
        self.clear_main()

        frame = ctk.CTkFrame(self.main)
        frame.pack(fill="both", expand=True)
        self.current_frame = frame

        ctk.CTkLabel(
            frame,
            text="📄 PDF → Word (OCR Avançado)",
            font=ctk.CTkFont(size=22, weight="bold")
        ).pack(pady=(30, 10))

        ctk.CTkLabel(
            frame,
            text="Converte PDF em Word editável preservando o máximo do layout",
            font=ctk.CTkFont(size=14)
        ).pack(pady=10)

        ctk.CTkButton(
            frame,
            text="Selecionar PDFs",
            width=280,
            height=50,
            font=ctk.CTkFont(size=15, weight="bold"),
            command=self.run_pdf_to_word_ocr
        ).pack(pady=30)

#================================================================================================================================
            
if __name__ == '__main__':
    app = App()
    app.mainloop()
